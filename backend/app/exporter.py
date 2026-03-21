from __future__ import annotations

import base64
from dataclasses import dataclass
import io
import json
import os
from pathlib import Path
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from html import escape
from typing import List, Sequence, Tuple
from urllib.request import Request, urlopen
from uuid import uuid4
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Mm, Pt, RGBColor

from .models import FormattingOptions, GenerateSecurityPayload, ThemePayload
from .parser import Node, render_preview_html, to_markdown, to_plain_text
from .security import (
    disable_docx_editing,
    mask_sensitive_url,
    remove_docx_metadata,
    secure_pdf,
    validate_remote_media_url,
)
from .themes import css_from_theme, watermark_html


@dataclass
class ExportResult:
    file_id: str
    output_path: Path
    warnings: List[str]
    requested_format: str
    actual_format: str
    conversion_engine: str = "native"
    external_fallback_used: bool = False
    warning: str | None = None


@dataclass
class CaptionTracker:
    chapter_idx: int = 0
    figure_global: int = 0
    table_global: int = 0
    figure_chapter: int = 0
    table_chapter: int = 0


def _env_flag(name: str, default: bool = False) -> bool:
    raw = os.environ.get(name)
    if raw is None:
        return default
    return str(raw).strip().lower() in {"1", "true", "yes", "on"}


def _allow_low_fidelity_pdf_fallback() -> bool:
    return _env_flag(
        "NF_PDF_ALLOW_LOW_FIDELITY_FALLBACK",
        default=True,
    )


def _docx2pdf_supported() -> bool:
    return sys.platform in {"win32", "darwin"}


def _hex_to_rgb(value: str) -> RGBColor:
    cleaned = value.strip().lstrip("#")
    if len(cleaned) != 6:
        return RGBColor(31, 58, 95)
    try:
        return RGBColor(int(cleaned[0:2], 16), int(cleaned[2:4], 16), int(cleaned[4:6], 16))
    except ValueError:
        return RGBColor(31, 58, 95)


def _mm_or_default(value: float, default: float) -> Mm:
    return Mm(value if value >= 0 else default)


def _font_primary(font_family: str) -> str:
    return (font_family.split(",")[0] if font_family else "Calibri").strip()


def _docx_alignment(value: str) -> WD_ALIGN_PARAGRAPH:
    if value == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if value == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if value == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def _style_num(styles: dict[str, object], *keys: str, default: float) -> float:
    for key in keys:
        if key not in styles:
            continue
        raw = styles.get(key)
        try:
            return float(raw)  # type: ignore[arg-type]
        except (TypeError, ValueError):
            continue
    return default


def _style_bool(styles: dict[str, object], *keys: str, default: bool) -> bool:
    for key in keys:
        if key not in styles:
            continue
        raw = styles.get(key)
        if isinstance(raw, bool):
            return raw
        if isinstance(raw, (int, float)):
            return bool(raw)
        if isinstance(raw, str):
            v = raw.strip().lower()
            if v in {"1", "true", "yes", "on"}:
                return True
            if v in {"0", "false", "no", "off"}:
                return False
    return default


def _style_str(styles: dict[str, object], *keys: str, default: str) -> str:
    for key in keys:
        if key not in styles:
            continue
        raw = styles.get(key)
        if raw is None:
            continue
        return str(raw).strip()
    return default


def _normalize_line_spacing(value: float) -> float:
    return max(1.0, min(3.0, round(float(value), 2)))


def _next_caption_number(state: CaptionTracker, kind: str) -> str:
    if kind == "figure":
        state.figure_global += 1
        if state.chapter_idx > 0:
            state.figure_chapter += 1
            return f"{state.chapter_idx}.{state.figure_chapter}"
        return str(state.figure_global)
    state.table_global += 1
    if state.chapter_idx > 0:
        state.table_chapter += 1
        return f"{state.chapter_idx}.{state.table_chapter}"
    return str(state.table_global)


def _collect_caption_entries(nodes: Sequence[Node]) -> tuple[list[str], list[str]]:
    state = CaptionTracker()
    figure_entries: list[str] = []
    table_entries: list[str] = []
    for node in nodes:
        if node.type == "chapter":
            state.chapter_idx += 1
            state.figure_chapter = 0
            state.table_chapter = 0
            continue
        if node.type in {"figure", "image"}:
            caption = (node.caption or node.text or "").strip()
            if node.type == "figure" or caption:
                number = _next_caption_number(state, "figure")
                figure_entries.append(f"Figure {number}: {caption or node.source or 'Image'}")
            continue
        if node.type == "figure_caption":
            number = _next_caption_number(state, "figure")
            figure_entries.append(f"Figure {number}: {node.text}")
            continue
        if node.type == "table_caption":
            number = _next_caption_number(state, "table")
            table_entries.append(f"Table {number}: {node.text}")
    return figure_entries, table_entries


def _append_toc_field(paragraph) -> None:
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)
    paragraph.add_run("Update this field in Word to generate the TOC.")

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def _decode_data_uri_image(source: str) -> bytes | None:
    match = re.match(r"^data:image/[a-zA-Z0-9.+-]+;base64,(.+)$", source, flags=re.IGNORECASE | re.DOTALL)
    if not match:
        return None
    try:
        return base64.b64decode(match.group(1), validate=False)
    except Exception:
        return None


def _resolve_image_stream(source: str, warnings: List[str]) -> io.BytesIO | None:
    value = (source or "").strip()
    if not value:
        return None
    try:
        decoded = _decode_data_uri_image(value)
        if decoded:
            return io.BytesIO(decoded)

        if value.lower().startswith(("http://", "https://")):
            ok, reason = validate_remote_media_url(
                value,
                allow_private=_env_flag("NF_ALLOW_PRIVATE_MEDIA_URLS", default=False),
            )
            if not ok:
                warnings.append(f"Blocked image URL '{mask_sensitive_url(value)}': {reason}")
                return None
            request = Request(
                value,
                method="GET",
                headers={"User-Agent": "quick-doc-formatter/8.0"},
            )
            with urlopen(request, timeout=12) as response:  # nosec B310
                payload = response.read()
                return io.BytesIO(payload)

        local = Path(value)
        if local.exists() and local.is_file():
            return io.BytesIO(local.read_bytes())
    except Exception as exc:
        warnings.append(f"Unable to load image source '{value}': {exc}")
        return None
    warnings.append(f"Image source not found: {value}")
    return None


def _apply_docx_background_watermark(section, security: GenerateSecurityPayload, theme: ThemePayload, warnings: List[str]) -> bool:
    watermark = security.watermark
    if not watermark or not watermark.value:
        return False
    if watermark.type != "text":
        return False

    try:
        header = section.header
        para = header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        rotation = int(float(watermark.rotation or 315))
        opacity = max(0.03, min(1.0, float(watermark.opacity or 0.10)))
        font_family = _font_primary(watermark.fontFamily or theme.fontFamily)
        font_size = max(28.0, float(watermark.size or 42.0))
        color = str(watermark.color or theme.primaryColor or "#8A8A8A").strip().lstrip("#")
        if len(color) != 6:
            color = "8A8A8A"

        shape_xml = (
            f'<w:pict {nsdecls("w", "v", "o")}>'
            '<v:shape id="QuickDocFormatterWatermark" type="#_x0000_t136"'
            ' style="position:absolute;'
            "mso-position-horizontal:center;"
            "mso-position-vertical:center;"
            "width:468pt;height:117pt;"
            "z-index:-251654144;"
            f'rotation:{rotation};"'
            f' fillcolor="#{xml_escape(color)}" stroked="f">'
            f'<v:fill opacity="{opacity:.2f}"/>'
            "<v:stroke on=\"f\"/>"
            '<v:textpath on="t" fitshape="t"'
            f' style="font-family:{xml_escape(font_family)};font-size:{font_size:.0f}pt"'
            f' string="{xml_escape(str(watermark.value))}"/>'
            "</v:shape>"
            "</w:pict>"
        )
        run._r.append(parse_xml(shape_xml))
        return True
    except Exception as exc:
        warnings.append(f"Background watermark insertion failed; using compatibility fallback. Detail: {exc}")
        return False


def _apply_docx_centered_image_watermark(section, security: GenerateSecurityPayload, warnings: List[str]) -> bool:
    watermark = security.watermark
    if not watermark or watermark.type != "image" or not watermark.value:
        return False
    image_stream = _resolve_image_stream(watermark.value, warnings)
    if not image_stream:
        return False
    try:
        para = section.header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        scale = max(10.0, min(120.0, float(watermark.scale or 38.0)))
        width_mm = max(20.0, min(170.0, 170.0 * (scale / 100.0)))
        para.add_run().add_picture(image_stream, width=Mm(width_mm))
        return True
    except Exception as exc:
        warnings.append(f"Image watermark insertion failed; skipped watermark image. Detail: {exc}")
        return False


def _field_instruction(field_name: str, number_style: str) -> str:
    style = (number_style or "arabic").strip().lower()
    suffix_map = {
        "arabic": "",
        "roman": r" \* ROMAN",
        "roman_lower": r" \* roman",
        "alpha": r" \* ALPHABETIC",
        "alpha_lower": r" \* alphabetic",
    }
    return f"{field_name}{suffix_map.get(style, '')}"


def _append_field(paragraph, field_name: str, number_style: str) -> None:
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = _field_instruction(field_name, number_style)
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    paragraph.add_run("1")

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def _normalize_page_template(mode: str, template: str) -> str:
    raw = (template or "").strip()
    if not raw:
        return "Page X of Y" if mode == "page_x_of_y" else "Page X"
    lowered = raw.lower()
    if "{page}" in lowered:
        raw = re.sub(r"\{page\}", "X", raw, flags=re.IGNORECASE)
    if "{pages}" in lowered or "{total}" in lowered:
        raw = re.sub(r"\{pages\}", "Y", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\{total\}", "Y", raw, flags=re.IGNORECASE)
    if "x" not in raw.lower() and "y" not in raw.lower():
        return f"{raw} X"
    if mode == "page_x_of_y" and "y" not in raw.lower():
        if "x" in raw.lower():
            return f"{raw} of Y"
        return "Page X of Y"
    return raw


def _add_page_number(
    paragraph,
    mode: str,
    *,
    number_style: str = "arabic",
    format_template: str = "",
) -> None:
    template = _normalize_page_template(mode, format_template)
    parts = re.split(r"(X|Y)", template, flags=re.IGNORECASE)
    for part in parts:
        if not part:
            continue
        token = part.upper()
        if token == "X":
            _append_field(paragraph, "PAGE", number_style)
        elif token == "Y":
            _append_field(paragraph, "NUMPAGES", number_style)
        else:
            paragraph.add_run(part)


def _docx_border_style(raw: str) -> str:
    value = (raw or "").strip().lower()
    mapping = {
        "single": "single",
        "double": "double",
        "dashed": "dashed",
        "dotted": "dotted",
        "thick": "thick",
    }
    return mapping.get(value, "single")


def _apply_paragraph_separator(paragraph, side: str, color: str, width_pt: float, style: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    edge = p_bdr.find(qn(f"w:{side}"))
    if edge is None:
        edge = OxmlElement(f"w:{side}")
        p_bdr.append(edge)

    color_hex = (color or "").strip().lstrip("#")
    if len(color_hex) != 6:
        color_hex = "BFBFBF"
    sz = max(2, min(96, int(round(max(0.25, width_pt) * 8))))
    edge.set(qn("w:val"), _docx_border_style(style))
    edge.set(qn("w:sz"), str(sz))
    edge.set(qn("w:space"), "1")
    edge.set(qn("w:color"), color_hex)


def _apply_page_borders(
    section,
    *,
    enabled: bool,
    width_pt: float,
    color: str,
    style: str,
    offset_pt: float,
) -> None:
    if not enabled:
        return

    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)

    pg_borders.set(qn("w:offsetFrom"), "page")
    color_hex = (color or "").strip().lstrip("#")
    if len(color_hex) != 6:
        color_hex = "000000"
    sz = max(2, min(96, int(round(max(0.25, width_pt) * 8))))
    spacing = max(0, min(96, int(round(max(0.0, offset_pt)))))
    border_style = _docx_border_style(style)

    for side in ("top", "left", "bottom", "right"):
        edge = pg_borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            pg_borders.append(edge)
        edge.set(qn("w:val"), border_style)
        edge.set(qn("w:sz"), str(sz))
        edge.set(qn("w:space"), str(spacing))
        edge.set(qn("w:color"), color_hex)


def _set_cell_borders(cell, *, color: str, width_pt: float, style: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    color_hex = (color or "").strip().lstrip("#")
    if len(color_hex) != 6:
        color_hex = "D1D5DB"
    sz = max(2, min(96, int(round(max(0.25, width_pt) * 8))))
    border_style = _docx_border_style(style)
    for side in ("top", "left", "bottom", "right"):
        edge = tc_borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            tc_borders.append(edge)
        edge.set(qn("w:val"), border_style)
        edge.set(qn("w:sz"), str(sz))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color_hex)


def _set_cell_shading(cell, fill_hex: str) -> None:
    cleaned = (fill_hex or "").strip().lstrip("#")
    if len(cleaned) != 6:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cleaned)


def _set_paragraph_shading(paragraph, fill_hex: str) -> None:
    cleaned = (fill_hex or "").strip().lstrip("#")
    if len(cleaned) != 6:
        return
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cleaned)


def _style_paragraph_runs(
    paragraph,
    *,
    font_name: str,
    size_pt: float,
    color_hex: str,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    color = _hex_to_rgb(color_hex)
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(max(1.0, size_pt))
        run.font.color.rgb = color
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic


def _inject_preview_running_blocks(html: str, security: GenerateSecurityPayload) -> str:
    header_text = (security.headerText or "").strip()
    footer_text = (security.footerText or "").strip()
    show_page_number = security.pageNumberMode in {"page_x", "page_x_of_y"}
    mode = security.pageNumberMode if show_page_number else "page_x"

    injected = []
    if header_text:
        injected.append(f'<div class="nf-running-header">{escape(header_text)}</div>')

    footer_parts = []
    if footer_text:
        footer_parts.append(f'<span class="nf-footer-text">{escape(footer_text)}</span>')
    if show_page_number:
        footer_parts.append(f'<span class="nf-page-num" data-mode="{mode}"></span>')
    if footer_parts:
        injected.append(f'<div class="nf-running-footer">{" ".join(footer_parts)}</div>')

    if not injected:
        return html
    return html.replace(
        '<div class="nf-preview-root">',
        f'<div class="nf-preview-root">{"".join(injected)}',
        1,
    )


def _find_libreoffice_binary() -> str | None:
    for name in ("soffice", "libreoffice"):
        path = shutil.which(name)
        if path:
            return path

    if os.name != "nt":
        candidates = [
            Path("/usr/bin/soffice"),
            Path("/usr/local/bin/soffice"),
            Path("/usr/lib/libreoffice/program/soffice"),
            Path("/opt/libreoffice/program/soffice"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)

    if os.name == "nt":
        candidates = [
            Path(os.environ.get("ProgramFiles", "")) / "LibreOffice" / "program" / "soffice.exe",
            Path(os.environ.get("ProgramFiles(x86)", "")) / "LibreOffice" / "program" / "soffice.exe",
            Path("C:/Program Files/LibreOffice/program/soffice.exe"),
            Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)
    return None


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> Tuple[bool, str]:
    errors: List[str] = []

    if _docx2pdf_supported():
        try:
            from docx2pdf import convert as docx2pdf_convert  # type: ignore

            docx2pdf_convert(str(docx_path), str(pdf_path))
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                return True, "docx2pdf"
            errors.append("docx2pdf produced no file")
        except Exception as exc:
            errors.append(f"docx2pdf failed: {exc}")
    else:
        errors.append(f"docx2pdf skipped on {sys.platform}")

    binary = _find_libreoffice_binary()
    if binary:
        cmd = [
            binary,
            "--headless",
            "--convert-to",
            "pdf",
            str(docx_path),
            "--outdir",
            str(pdf_path.parent),
        ]
        try:
            result = subprocess.run(cmd, timeout=90, check=False, capture_output=True, text=True)
            produced = pdf_path.parent / f"{docx_path.stem}.pdf"
            if result.returncode == 0 and produced.exists() and produced.stat().st_size > 0:
                if produced != pdf_path:
                    produced.replace(pdf_path)
                return True, "libreoffice"
            errors.append(result.stderr.strip() or result.stdout.strip() or "libreoffice failed")
        except Exception as exc:
            errors.append(f"libreoffice exception: {exc}")
    else:
        errors.append("libreoffice binary not found")

    return False, "; ".join(err for err in errors if err)


def _http_json(
    method: str,
    url: str,
    *,
    token: str | None = None,
    payload: dict | None = None,
    timeout: int = 60,
) -> dict:
    headers = {"Accept": "application/json", "User-Agent": "quick-doc-formatter/8.0"}
    body = None
    if token:
        headers["Authorization"] = f"Bearer {token}"
    if payload is not None:
        headers["Content-Type"] = "application/json"
        body = json.dumps(payload).encode("utf-8")
    req = Request(url=url, data=body, method=method.upper(), headers=headers)
    with urlopen(req, timeout=timeout) as resp:  # nosec B310
        raw = resp.read()
    return json.loads(raw.decode("utf-8"))


def _encode_multipart_form(fields: dict[str, str], file_field: str, file_name: str, file_bytes: bytes) -> tuple[bytes, str]:
    boundary = f"----QuickDocFormatterBoundary{uuid4().hex}"
    chunks: list[bytes] = []
    for key, value in fields.items():
        chunks.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{key}"\r\n\r\n'.encode("utf-8"),
                str(value).encode("utf-8"),
                b"\r\n",
            ]
        )
    chunks.extend(
        [
            f"--{boundary}\r\n".encode("utf-8"),
            (
                f'Content-Disposition: form-data; name="{file_field}"; filename="{file_name}"\r\n'
                "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
            ).encode("utf-8"),
            file_bytes,
            b"\r\n",
            f"--{boundary}--\r\n".encode("utf-8"),
        ]
    )
    return b"".join(chunks), f"multipart/form-data; boundary={boundary}"


def _convert_docx_to_pdf_ilovepdf_api(docx_path: Path, pdf_path: Path) -> Tuple[bool, str]:
    public_key = (os.environ.get("NF_ILOVEPDF_PUBLIC_KEY") or "").strip()
    if not public_key:
        return False, "ilovepdf api not configured"

    auth_url = os.environ.get("NF_ILOVEPDF_AUTH_URL", "https://api.ilovepdf.com/v1/auth").strip()
    start_base = os.environ.get("NF_ILOVEPDF_START_URL", "https://api.ilovepdf.com/v1/start").strip().rstrip("/")
    region = (os.environ.get("NF_ILOVEPDF_REGION") or "us").strip().lower() or "us"
    if region not in {"eu", "us", "fr", "de", "pl"}:
        region = "us"

    try:
        token_payload = _http_json("POST", auth_url, payload={"public_key": public_key}, timeout=30)
        token = str(token_payload.get("token") or "").strip()
        if not token:
            return False, "ilovepdf auth returned no token"

        start_payload = _http_json("GET", f"{start_base}/officepdf/{region}", token=token, timeout=30)
        server = str(start_payload.get("server") or "").strip()
        task = str(start_payload.get("task") or "").strip()
        if not server or not task:
            return False, "ilovepdf start did not return server/task"

        upload_url = f"https://{server}/v1/upload"
        body, content_type = _encode_multipart_form(
            {"task": task},
            "file",
            docx_path.name,
            docx_path.read_bytes(),
        )
        upload_headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": content_type,
            "Accept": "application/json",
            "User-Agent": "quick-doc-formatter/8.0",
        }
        upload_req = Request(upload_url, data=body, method="POST", headers=upload_headers)
        with urlopen(upload_req, timeout=90) as upload_resp:  # nosec B310
            upload_payload = json.loads(upload_resp.read().decode("utf-8"))
        server_filename = str(upload_payload.get("server_filename") or "").strip()
        if not server_filename:
            return False, "ilovepdf upload did not return server_filename"

        process_url = f"https://{server}/v1/process"
        process_payload = {
            "task": task,
            "tool": "officepdf",
            "output_filename": docx_path.stem,
            "ignore_errors": True,
            "files": [
                {
                    "server_filename": server_filename,
                    "filename": docx_path.name,
                }
            ],
        }
        _http_json("POST", process_url, token=token, payload=process_payload, timeout=180)

        download_url = f"https://{server}/v1/download/{task}"
        dl_req = Request(
            url=download_url,
            method="GET",
            headers={
                "Authorization": f"Bearer {token}",
                "Accept": "application/pdf,application/zip,application/octet-stream",
                "User-Agent": "quick-doc-formatter/8.0",
            },
        )
        with urlopen(dl_req, timeout=180) as dl_resp:  # nosec B310
            pdf_bytes = dl_resp.read()

        if not pdf_bytes:
            return False, "ilovepdf download returned empty payload"
        if pdf_bytes.startswith(b"PK"):
            try:
                with zipfile.ZipFile(io.BytesIO(pdf_bytes)) as archive:
                    pdf_members = [name for name in archive.namelist() if name.lower().endswith(".pdf")]
                    if not pdf_members:
                        return False, "ilovepdf zip download did not contain a pdf file"
                    pdf_bytes = archive.read(pdf_members[0])
            except Exception as exc:
                return False, f"ilovepdf zip extraction failed: {exc}"
        pdf_path.write_bytes(pdf_bytes)
        if not pdf_path.exists() or pdf_path.stat().st_size == 0:
            return False, "ilovepdf produced empty output file"
        return True, "ilovepdf_api"
    except Exception as exc:
        return False, f"ilovepdf api failed: {exc}"


def _convert_docx_to_pdf_ilovepdf_automation(docx_path: Path, pdf_path: Path) -> Tuple[bool, str]:
    automation_url = (os.environ.get("NF_ILOVEPDF_AUTOMATION_URL") or "").strip()
    if not automation_url:
        return False, "ilovepdf automation bridge not configured"

    ok, reason = validate_remote_media_url(
        automation_url,
        allow_private=_env_flag("NF_ALLOW_PRIVATE_MEDIA_URLS", default=False),
    )
    if not ok:
        return False, f"ilovepdf automation url blocked: {reason}"

    try:
        body, content_type = _encode_multipart_form(
            {"filename": docx_path.name},
            "file",
            docx_path.name,
            docx_path.read_bytes(),
        )
        req = Request(
            url=automation_url,
            data=body,
            method="POST",
            headers={
                "Content-Type": content_type,
                "Accept": "application/pdf,application/octet-stream",
                "User-Agent": "quick-doc-formatter/8.0",
            },
        )
        with urlopen(req, timeout=240) as resp:  # nosec B310
            payload = resp.read()
        if not payload:
            return False, "ilovepdf automation returned empty payload"
        pdf_path.write_bytes(payload)
        if not pdf_path.exists() or pdf_path.stat().st_size == 0:
            return False, "ilovepdf automation produced empty output file"
        return True, "ilovepdf_automation"
    except Exception as exc:
        return False, f"ilovepdf automation failed: {exc}"


def _convert_html_to_pdf_weasyprint(html: str, pdf_path: Path) -> Tuple[bool, str]:
    try:
        from weasyprint import HTML  # type: ignore

        HTML(string=html).write_pdf(str(pdf_path))
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return True, "weasyprint"
        return False, "weasyprint produced no file"
    except Exception as exc:
        return False, f"weasyprint failed: {exc}"


def _nodes_to_plain_lines(nodes: Sequence[Node]) -> list[str]:
    lines: list[str] = []
    for node in nodes:
        if node.type == "heading":
            lines.append(node.text.strip())
            lines.append("")
        elif node.type in {"section", "chapter", "appendix", "references_heading"}:
            lines.append(node.text.strip())
            lines.append("")
        elif node.type in {"toc", "list_of_tables", "list_of_figures"}:
            lines.append((node.text or node.type.replace("_", " ").title()).strip())
            lines.append("")
        elif node.type == "reference":
            lines.append(f"- {node.text}".strip())
            lines.append("")
        elif node.type == "paragraph":
            lines.append(node.text.strip())
            lines.append("")
        elif node.type == "bullet":
            for item in node.items or []:
                lines.append(f"- {item}".strip())
            lines.append("")
        elif node.type == "numbered":
            for idx, item in enumerate(node.items or [], start=1):
                lines.append(f"{idx}. {item}".strip())
            lines.append("")
        elif node.type == "checklist":
            checks = node.checks or []
            for idx, item in enumerate(node.items or []):
                checked = checks[idx] if idx < len(checks) else False
                lines.append(f"[{'x' if checked else ' '}] {item}".strip())
            lines.append("")
        elif node.type == "code":
            lines.append("CODE:")
            for ln in node.text.splitlines() or [""]:
                lines.append(f"  {ln}")
            lines.append("")
        elif node.type == "equation":
            lines.append("EQUATION:")
            for ln in node.text.splitlines() or [""]:
                lines.append(f"  {ln}")
            lines.append("")
        elif node.type == "ascii":
            lines.append(node.text)
        elif node.type == "table":
            for row in node.rows or []:
                lines.append(" | ".join(row))
            lines.append("")
        elif node.type in {"image", "figure"}:
            lines.append(f"[IMAGE] {node.source}".strip())
            if node.caption:
                lines.append(node.caption)
            lines.append("")
        elif node.type in {"table_caption", "figure_caption"}:
            lines.append(node.text.strip())
            lines.append("")
        elif node.type == "separator":
            lines.append("--------------------")
            lines.append("")
        elif node.type == "pagebreak":
            lines.append("__NF_PAGE_BREAK__")
    return lines or ["(empty document)"]


def _convert_nodes_to_pdf_reportlab(
    nodes: Sequence[Node],
    theme: ThemePayload,
    formatting: FormattingOptions,
    security: GenerateSecurityPayload,
    pdf_path: Path,
) -> Tuple[bool, str]:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
    except Exception as exc:
        return False, f"reportlab unavailable: {exc}"

    try:
        lines = _nodes_to_plain_lines(nodes)
        page_w, page_h = A4
        margin_top = (formatting.margins.top or 25) * mm
        margin_bottom = (formatting.margins.bottom or 25) * mm
        margin_left = (formatting.margins.left or 25) * mm
        margin_right = (formatting.margins.right or 25) * mm
        usable_h = max(60.0, page_h - margin_top - margin_bottom)

        font_size = float(theme.bodyStyle.size or 12)
        line_height = max(12.0, font_size * 1.55)
        max_lines = max(5, int(usable_h // line_height) - 2)

        pages: list[list[str]] = []
        cur: list[str] = []
        for ln in lines:
            if ln == "__NF_PAGE_BREAK__":
                pages.append(cur)
                cur = []
                continue
            cur.append(ln)
            if len(cur) >= max_lines:
                pages.append(cur)
                cur = []
        if cur or not pages:
            pages.append(cur)

        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        c.setAuthor("Quick Doc Formatter")
        c.setTitle("Quick Doc Formatter Export")
        try:
            c.setSubject(theme.name or "Quick Doc Formatter Document")
        except Exception:
            pass

        for p_idx, page_lines in enumerate(pages, start=1):
            y = page_h - margin_top
            if security.headerText:
                c.setFont("Helvetica-Bold", 10)
                c.drawString(margin_left, y, str(security.headerText))
                y -= line_height

            c.setFont("Helvetica", font_size)
            for ln in page_lines:
                text = (ln or "").replace("\t", "    ")
                max_chars = max(20, int((page_w - margin_left - margin_right) / (font_size * 0.55)))
                while len(text) > max_chars:
                    c.drawString(margin_left, y, text[:max_chars])
                    text = text[max_chars:]
                    y -= line_height
                    if y < margin_bottom + line_height:
                        break
                if y < margin_bottom + line_height:
                    break
                c.drawString(margin_left, y, text)
                y -= line_height

            footer_parts: list[str] = []
            if security.footerText:
                footer_parts.append(str(security.footerText))
            mode = security.pageNumberMode or "page_x"
            if mode == "page_x_of_y":
                footer_parts.append(f"Page {p_idx} of {len(pages)}")
            else:
                footer_parts.append(f"Page {p_idx}")
            c.setFont("Helvetica", 9)
            c.drawString(margin_left, margin_bottom * 0.55, " | ".join(footer_parts))
            c.showPage()

        c.save()
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return True, "reportlab"
        return False, "reportlab produced no file"
    except Exception as exc:
        return False, f"reportlab failed: {exc}"


def _pdf_escape_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_emergency_pdf(pdf_path: Path, message: str) -> Tuple[bool, str]:
    try:
        line = _pdf_escape_text("Quick Doc Formatter PDF fallback: " + (message or "Content generated."))
        stream = f"BT /F1 11 Tf 50 780 Td ({line}) Tj ET".encode("latin-1", "replace")

        objects: list[bytes] = []
        objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
        objects.append(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
        objects.append(
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>\nendobj\n"
        )
        objects.append(
            b"4 0 obj\n<< /Length "
            + str(len(stream)).encode("ascii")
            + b" >>\nstream\n"
            + stream
            + b"\nendstream\nendobj\n"
        )
        objects.append(b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n")

        pdf = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets = [0]
        for obj in objects:
            offsets.append(len(pdf))
            pdf.extend(obj)
        xref_start = len(pdf)
        pdf.extend(f"xref\n0 {len(offsets)}\n".encode("ascii"))
        pdf.extend(b"0000000000 65535 f \n")
        for off in offsets[1:]:
            pdf.extend(f"{off:010d} 00000 n \n".encode("ascii"))
        pdf.extend(
            (
                "trailer\n<< /Size {size} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".format(
                    size=len(offsets),
                    xref=xref_start,
                )
            ).encode("ascii")
        )
        pdf_path.write_bytes(bytes(pdf))
        return True, "emergency_pdf"
    except Exception as exc:
        return False, f"emergency_pdf failed: {exc}"


class FileStore:
    def __init__(self, base_dir: str, ttl_seconds: int = 60 * 60 * 6) -> None:
        root = Path(base_dir) if base_dir else Path(tempfile.gettempdir()) / "quick_doc_formatter_exports"
        root.mkdir(parents=True, exist_ok=True)
        self.base_dir = root
        self.ttl_seconds = ttl_seconds

    def cleanup(self) -> None:
        now = time.time()
        for file_path in self.base_dir.glob("*"):
            if not file_path.is_file():
                continue
            age = now - file_path.stat().st_mtime
            if age > self.ttl_seconds:
                try:
                    file_path.unlink()
                except OSError:
                    continue

    def reserve_path(self, ext: str) -> Tuple[str, Path]:
        file_id = uuid4().hex
        path = self.base_dir / f"{file_id}.{ext}"
        return file_id, path

    def resolve_path(self, file_id: str) -> Path | None:
        if not file_id or not all(ch in "0123456789abcdef" for ch in file_id.lower()):
            return None
        matches = list(self.base_dir.glob(f"{file_id}.*"))
        return matches[0] if matches else None


class DocumentExporter:
    def __init__(self, store: FileStore) -> None:
        self.store = store

    def create_preview_html(
        self,
        nodes: Sequence[Node],
        theme: ThemePayload,
        formatting: FormattingOptions,
        security: GenerateSecurityPayload,
        include_running_blocks: bool = True,
    ) -> str:
        css = css_from_theme(theme, formatting)
        watermark = watermark_html(security.watermark)
        html = render_preview_html(nodes, css, watermark)
        if include_running_blocks:
            return _inject_preview_running_blocks(html, security)
        return html

    def create_docx(
        self,
        nodes: Sequence[Node],
        theme: ThemePayload,
        formatting: FormattingOptions,
        security: GenerateSecurityPayload,
    ) -> Tuple[str, Path, List[str]]:
        self.store.cleanup()
        warnings: List[str] = []
        file_id, path = self.store.reserve_path("docx")

        document = Document()
        section = document.sections[0]
        styles = theme.styles if isinstance(theme.styles, dict) else {}

        page_size = _style_str(styles, "page_size", "pageSize", default="A4").upper()
        page_orientation = _style_str(
            styles,
            "page_orientation",
            "pageOrientation",
            default="portrait",
        ).lower()
        if page_size in {"A4", "A3", "LETTER", "LEGAL"}:
            page_sizes_mm = {
                "A4": (210.0, 297.0),
                "A3": (297.0, 420.0),
                "LETTER": (215.9, 279.4),
                "LEGAL": (215.9, 355.6),
            }
            w_mm, h_mm = page_sizes_mm[page_size]
            if page_orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Mm(h_mm)
                section.page_height = Mm(w_mm)
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = Mm(w_mm)
                section.page_height = Mm(h_mm)

        section.top_margin = _mm_or_default(formatting.margins.top, 25)
        section.bottom_margin = _mm_or_default(formatting.margins.bottom, 25)
        section.left_margin = _mm_or_default(formatting.margins.left, 25)
        section.right_margin = _mm_or_default(formatting.margins.right, 25)

        font_name = _font_primary(theme.fontFamily)
        body_color = _style_str(styles, "body_color", "bodyColor", default="#17202a")
        code_font_name = _font_primary(
            _style_str(
                styles,
                "code_font_family",
                "codeFontFamily",
                default="JetBrains Mono, Consolas, Courier New, monospace",
            )
        )
        bullet_font_name = _font_primary(
            _style_str(styles, "bullet_font_family", "bulletFontFamily", default=font_name)
        )
        code_font_size = _style_num(styles, "code_font_size", "codeFontSize", default=10.0)
        line_spacing = _normalize_line_spacing(formatting.lineSpacing or theme.bodyStyle.lineHeight or 1.5)

        normal_style = document.styles["Normal"]
        normal_style.font.name = font_name
        normal_style.font.size = Pt(theme.bodyStyle.size or 12)
        normal_style.font.color.rgb = _hex_to_rgb(body_color)
        normal_style.paragraph_format.line_spacing = line_spacing

        _apply_page_borders(
            section,
            enabled=_style_bool(styles, "page_border_enabled", "pageBorderEnabled", default=False),
            width_pt=_style_num(
                styles,
                "page_border_width",
                "pageBorderWidth",
                default=1.0,
            ),
            color=_style_str(
                styles,
                "page_border_color",
                "pageBorderColor",
                default=theme.primaryColor,
            ),
            style=_style_str(
                styles,
                "page_border_style",
                "pageBorderStyle",
                default="single",
            ),
            offset_pt=_style_num(
                styles,
                "page_border_offset",
                "pageBorderOffset",
                default=24.0,
            ),
        )

        header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        footer_para = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        header_para.text = (security.headerText or "").strip()
        footer_para.text = (security.footerText or "").strip()

        header_align = _style_str(styles, "header_alignment", "headerAlignment", default="center").lower()
        footer_align = _style_str(styles, "footer_alignment", "footerAlignment", default="center").lower()
        header_para.alignment = _docx_alignment(header_align)
        footer_para.alignment = _docx_alignment(footer_align)

        header_size = _style_num(styles, "header_size", "headerSize", default=10.0)
        footer_size = _style_num(styles, "footer_size", "footerSize", default=9.0)
        header_color = _style_str(styles, "header_color", "headerColor", default=theme.primaryColor)
        footer_color = _style_str(styles, "footer_color", "footerColor", default=theme.primaryColor)
        header_font = _font_primary(
            _style_str(styles, "header_font_family", "headerFontFamily", default=font_name)
        )
        footer_font = _font_primary(
            _style_str(styles, "footer_font_family", "footerFontFamily", default=font_name)
        )
        header_bold = _style_bool(styles, "header_bold", "headerBold", default=False)
        footer_bold = _style_bool(styles, "footer_bold", "footerBold", default=False)
        header_italic = _style_bool(styles, "header_italic", "headerItalic", default=False)
        footer_italic = _style_bool(styles, "footer_italic", "footerItalic", default=False)
        _style_paragraph_runs(
            header_para,
            font_name=header_font,
            size_pt=header_size,
            color_hex=header_color,
            bold=header_bold,
            italic=header_italic,
        )
        _style_paragraph_runs(
            footer_para,
            font_name=footer_font,
            size_pt=footer_size,
            color_hex=footer_color,
            bold=footer_bold,
            italic=footer_italic,
        )

        if _style_bool(styles, "header_separator", "headerSeparator", default=False):
            _apply_paragraph_separator(
                header_para,
                "bottom",
                _style_str(styles, "header_separator_color", "headerSeparatorColor", default="#CFCFCF"),
                0.75,
                "single",
            )
        if _style_bool(styles, "footer_separator", "footerSeparator", default=False):
            _apply_paragraph_separator(
                footer_para,
                "top",
                _style_str(styles, "footer_separator_color", "footerSeparatorColor", default="#CFCFCF"),
                0.75,
                "single",
            )

        page_mode = security.pageNumberMode or _style_str(
            styles,
            "page_number_mode",
            "pageNumberMode",
            default="page_x",
        )
        if page_mode not in {"page_x", "page_x_of_y"}:
            page_mode = "page_x"
        page_number_position = _style_str(
            styles,
            "page_number_position",
            "pageNumberPosition",
            default="footer",
        ).lower()
        if page_number_position not in {"header", "footer"}:
            page_number_position = "footer"
        show_header_page_numbers = _style_bool(
            styles,
            "header_show_page_numbers",
            "headerShowPageNumbers",
            default=False,
        )
        show_footer_page_numbers = _style_bool(
            styles,
            "footer_show_page_numbers",
            "footerShowPageNumbers",
            default=True,
        )
        page_number_alignment = _style_str(
            styles,
            "page_number_alignment",
            "pageNumberAlignment",
            default="",
        ).lower()
        if not page_number_alignment:
            page_number_alignment = (
                header_align if page_number_position == "header" else footer_align
            )
        page_number_style = _style_str(
            styles,
            "page_number_style",
            "pageNumberStyle",
            default="arabic",
        ).lower()
        if page_number_style not in {"arabic", "roman", "roman_lower", "alpha", "alpha_lower"}:
            page_number_style = "arabic"
        page_number_format = _style_str(
            styles,
            "page_number_format",
            "pageNumberFormat",
            default="Page X of Y" if page_mode == "page_x_of_y" else "Page X",
        )

        page_para = header_para if page_number_position == "header" else footer_para
        include_page_numbers = (
            show_header_page_numbers if page_number_position == "header" else show_footer_page_numbers
        )
        if include_page_numbers:
            if page_para.runs and page_para.runs[-1].text and not page_para.runs[-1].text.endswith(" "):
                page_para.add_run(" ")
            _add_page_number(
                page_para,
                page_mode,
                number_style=page_number_style,
                format_template=page_number_format,
            )
            page_para.alignment = _docx_alignment(page_number_alignment)
            is_header = page_number_position == "header"
            _style_paragraph_runs(
                page_para,
                font_name=header_font if is_header else footer_font,
                size_pt=header_size if is_header else footer_size,
                color_hex=header_color if is_header else footer_color,
                bold=header_bold if is_header else footer_bold,
                italic=header_italic if is_header else footer_italic,
            )

        paragraph_after_pt = _style_num(
            styles,
            "paragraph_spacing_after",
            "paragraphSpacingAfter",
            default=0.0,
        )
        paragraph_before_pt = _style_num(
            styles,
            "paragraph_spacing_before",
            "paragraphSpacingBefore",
            default=0.0,
        )
        heading_before_pt = _style_num(
            styles,
            "heading_spacing_before",
            "headingSpacingBefore",
            default=10.0,
        )
        heading_after_pt = _style_num(
            styles,
            "heading_spacing_after",
            "headingSpacingAfter",
            default=6.0,
        )
        paragraph_first_indent_em = _style_num(
            styles,
            "paragraph_first_line_indent",
            "paragraphFirstLineIndent",
            default=0.0,
        )
        bullet_base_indent_in = _style_num(
            styles,
            "bullet_base_indent",
            "bulletBaseIndent",
            default=0.25,
        )
        bullet_indent_per_level_in = _style_num(
            styles,
            "bullet_indent_per_level",
            "bulletIndentPerLevel",
            default=0.45,
        )
        code_indent_in = _style_num(
            styles,
            "code_indent",
            "codeIndent",
            default=0.0,
        )
        quote_indent_in = _style_num(
            styles,
            "quote_indent",
            "quoteIndent",
            default=0.5,
        )
        paragraph_alignment = _style_str(
            styles,
            "paragraph_alignment",
            "paragraphAlignment",
            default="left",
        ).lower()
        table_text_alignment = _style_str(
            styles,
            "table_text_alignment",
            "tableTextAlignment",
            default="left",
        ).lower()
        table_border_width = float(theme.tableStyle.borderWidth or 1)
        table_border_color = theme.tableStyle.borderColor or "#d1d5db"
        table_border_style = _style_str(
            styles,
            "table_border_style",
            "tableBorderStyle",
            default="single",
        )
        table_header_fill = theme.tableStyle.headerFill or "#f3f4f6"
        table_header_text = _style_str(
            styles,
            "table_header_text",
            "tableHeaderText",
            default="#111827",
        )
        table_odd_fill = _style_str(
            styles,
            "table_odd_row",
            "tableOddRow",
            default="#FFFFFF",
        )
        table_even_fill = _style_str(
            styles,
            "table_even_row",
            "tableEvenRow",
            default="#F8FAFC",
        )

        if security.watermark and security.watermark.value:
            if security.watermark.type == "text":
                applied_background_watermark = _apply_docx_background_watermark(section, security, theme, warnings)
                if not applied_background_watermark:
                    # Compatibility fallback: keep watermark centered in the header layer.
                    p = section.header.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(str(security.watermark.value))
                    run.font.name = _font_primary(security.watermark.fontFamily or theme.fontFamily)
                    run.font.color.rgb = _hex_to_rgb(security.watermark.color or theme.primaryColor)
                    run.font.size = Pt(max(20.0, float(security.watermark.size or 36)))
                    run.font.bold = True
            elif security.watermark.type == "image":
                if not _apply_docx_centered_image_watermark(section, security, warnings):
                    warnings.append("Image watermark source not found; skipped watermark image.")

        figure_entries, table_entries = _collect_caption_entries(nodes)
        caption_state = CaptionTracker()
        has_written_content = False

        for node in nodes:
            if node.type in {"section", "chapter", "appendix", "references_heading"}:
                if has_written_content:
                    document.add_page_break()
                if node.type == "chapter":
                    caption_state.chapter_idx += 1
                    caption_state.figure_chapter = 0
                    caption_state.table_chapter = 0
                    title = f"CHAPTER {caption_state.chapter_idx}: {node.text or 'Chapter'}"
                    p = document.add_heading(title, level=1)
                elif node.type == "appendix":
                    p = document.add_heading(f"Appendix: {node.text or 'Appendix'}", level=1)
                else:
                    p = document.add_heading(node.text or node.type.replace("_", " ").title(), level=1)
                p.paragraph_format.space_before = Pt(max(0.0, heading_before_pt))
                p.paragraph_format.space_after = Pt(max(0.0, heading_after_pt))
                p.paragraph_format.line_spacing = line_spacing
                if p.runs:
                    run = p.runs[0]
                    run.font.name = _font_primary(_style_str(styles, "h1_family", "h1Family", default=font_name))
                    run.font.size = Pt(getattr(theme.headingStyle, "h1").size or 18)
                    run.font.bold = True
                    run.font.color.rgb = _hex_to_rgb(getattr(theme.headingStyle, "h1").color or theme.primaryColor)
            elif node.type == "toc":
                p = document.add_heading(node.text or "Table of Contents", level=1)
                p.paragraph_format.line_spacing = line_spacing
                toc_p = document.add_paragraph()
                _append_toc_field(toc_p)
                toc_p.paragraph_format.line_spacing = line_spacing
                _style_paragraph_runs(
                    toc_p,
                    font_name=font_name,
                    size_pt=float(theme.bodyStyle.size or 12),
                    color_hex=body_color,
                )
            elif node.type == "list_of_tables":
                p = document.add_heading(node.text or "List of Tables", level=1)
                p.paragraph_format.line_spacing = line_spacing
                for entry in table_entries or ["Table entries are generated when captions are available."]:
                    item = document.add_paragraph(entry, style="List Bullet")
                    item.paragraph_format.line_spacing = line_spacing
                    _style_paragraph_runs(
                        item,
                        font_name=font_name,
                        size_pt=float(theme.bodyStyle.size or 12),
                        color_hex=body_color,
                    )
            elif node.type == "list_of_figures":
                p = document.add_heading(node.text or "List of Figures", level=1)
                p.paragraph_format.line_spacing = line_spacing
                for entry in figure_entries or ["Figure entries are generated when captions are available."]:
                    item = document.add_paragraph(entry, style="List Bullet")
                    item.paragraph_format.line_spacing = line_spacing
                    _style_paragraph_runs(
                        item,
                        font_name=font_name,
                        size_pt=float(theme.bodyStyle.size or 12),
                        color_hex=body_color,
                    )
            elif node.type == "reference":
                p = document.add_paragraph(node.text, style="List Number")
                p.paragraph_format.line_spacing = line_spacing
                _style_paragraph_runs(
                    p,
                    font_name=font_name,
                    size_pt=float(theme.bodyStyle.size or 12),
                    color_hex=body_color,
                )
            elif node.type == "heading":
                level = max(1, min(6, node.level))
                p = document.add_heading(node.text, level=level)
                p.paragraph_format.space_before = Pt(max(0.0, heading_before_pt))
                p.paragraph_format.space_after = Pt(max(0.0, heading_after_pt))
                p.paragraph_format.line_spacing = line_spacing
                if p.runs:
                    token = getattr(theme.headingStyle, f"h{level}")
                    run = p.runs[0]
                    run.font.name = _font_primary(
                        _style_str(
                            styles,
                            f"h{level}_family",
                            f"h{level}Family",
                            default=font_name,
                        )
                    )
                    run.font.size = Pt(token.size or max(12, 26 - (level * 2)))
                    run.font.bold = str(token.weight or "600") in {"600", "700", "800", "900"}
                    run.font.color.rgb = _hex_to_rgb(token.color or theme.primaryColor)
            elif node.type == "paragraph":
                p = document.add_paragraph(node.text)
                effective_align = node.align if node.align and node.align != "left" else paragraph_alignment
                p.alignment = _docx_alignment(effective_align)
                p.paragraph_format.space_before = Pt(max(0.0, paragraph_before_pt))
                p.paragraph_format.space_after = Pt(max(0.0, paragraph_after_pt))
                p.paragraph_format.line_spacing = line_spacing
                p.paragraph_format.first_line_indent = Inches(max(0.0, paragraph_first_indent_em) * 0.15)
                if getattr(node, "role", "paragraph") == "quote":
                    p.paragraph_format.left_indent = Inches(max(0.0, quote_indent_in))
                _style_paragraph_runs(
                    p,
                    font_name=font_name,
                    size_pt=float(theme.bodyStyle.size or 12),
                    color_hex=body_color,
                )
                role = getattr(node, "role", "paragraph")
                role_fill_map = {
                    "tip": "E0F2FE",
                    "warning": "FEF3C7",
                    "info": "DBEAFE",
                    "success": "DCFCE7",
                    "callout": "EEF2FF",
                    "summary": "F1F5F9",
                }
                if role in role_fill_map:
                    _set_paragraph_shading(p, role_fill_map[role])
            elif node.type == "bullet":
                levels = node.levels or []
                for idx, item in enumerate(node.items or []):
                    item_level = levels[idx] if idx < len(levels) else 0
                    p = document.add_paragraph(item, style="List Bullet")
                    p.alignment = _docx_alignment(paragraph_alignment)
                    p.paragraph_format.left_indent = Inches(
                        max(0.0, bullet_base_indent_in + (bullet_indent_per_level_in * item_level))
                    )
                    p.paragraph_format.space_before = Pt(max(0.0, paragraph_before_pt))
                    p.paragraph_format.space_after = Pt(max(0.0, paragraph_after_pt))
                    p.paragraph_format.line_spacing = line_spacing
                    _style_paragraph_runs(
                        p,
                        font_name=bullet_font_name,
                        size_pt=float(theme.bodyStyle.size or 12),
                        color_hex=body_color,
                    )
            elif node.type == "checklist":
                levels = node.levels or []
                checks = node.checks or []
                for idx, item in enumerate(node.items or []):
                    item_level = levels[idx] if idx < len(levels) else 0
                    checked = checks[idx] if idx < len(checks) else False
                    label = f"[{'x' if checked else ' '}] {item}"
                    p = document.add_paragraph(label, style="List Bullet")
                    p.alignment = _docx_alignment(paragraph_alignment)
                    p.paragraph_format.left_indent = Inches(
                        max(0.0, bullet_base_indent_in + (bullet_indent_per_level_in * item_level))
                    )
                    p.paragraph_format.space_before = Pt(max(0.0, paragraph_before_pt))
                    p.paragraph_format.space_after = Pt(max(0.0, paragraph_after_pt))
                    p.paragraph_format.line_spacing = line_spacing
                    _style_paragraph_runs(
                        p,
                        font_name=bullet_font_name,
                        size_pt=float(theme.bodyStyle.size or 12),
                        color_hex=body_color,
                    )
            elif node.type == "numbered":
                levels = node.levels or []
                for idx, item in enumerate(node.items or []):
                    item_level = levels[idx] if idx < len(levels) else 0
                    p = document.add_paragraph(item, style="List Number")
                    p.alignment = _docx_alignment(paragraph_alignment)
                    p.paragraph_format.left_indent = Inches(
                        max(0.0, bullet_base_indent_in + (bullet_indent_per_level_in * item_level))
                    )
                    p.paragraph_format.space_before = Pt(max(0.0, paragraph_before_pt))
                    p.paragraph_format.space_after = Pt(max(0.0, paragraph_after_pt))
                    p.paragraph_format.line_spacing = line_spacing
                    _style_paragraph_runs(
                        p,
                        font_name=bullet_font_name,
                        size_pt=float(theme.bodyStyle.size or 12),
                        color_hex=body_color,
                    )
            elif node.type == "code":
                p = document.add_paragraph(node.text)
                p.paragraph_format.left_indent = Inches(max(0.0, code_indent_in))
                p.paragraph_format.space_before = Pt(max(0.0, paragraph_before_pt))
                p.paragraph_format.space_after = Pt(max(0.0, paragraph_after_pt))
                p.paragraph_format.line_spacing = line_spacing
                _set_paragraph_shading(
                    p,
                    _style_str(styles, "code_background", "codeBackground", default="#0f172a"),
                )
                _style_paragraph_runs(
                    p,
                    font_name=code_font_name,
                    size_pt=code_font_size,
                    color_hex=_style_str(styles, "code_text", "codeText", default="#1f2937"),
                )
            elif node.type == "equation":
                p = document.add_paragraph(node.text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.left_indent = Inches(max(0.0, code_indent_in))
                p.paragraph_format.space_before = Pt(max(0.0, paragraph_before_pt))
                p.paragraph_format.space_after = Pt(max(0.0, paragraph_after_pt))
                p.paragraph_format.line_spacing = line_spacing
                _set_paragraph_shading(
                    p,
                    _style_str(styles, "code_background", "codeBackground", default="#0f172a"),
                )
                _style_paragraph_runs(
                    p,
                    font_name=code_font_name,
                    size_pt=code_font_size,
                    color_hex=_style_str(styles, "code_text", "codeText", default="#1f2937"),
                    italic=True,
                )
            elif node.type == "ascii":
                p = document.add_paragraph(node.text)
                p.paragraph_format.left_indent = Inches(0)
                p.paragraph_format.line_spacing = line_spacing
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_paragraph_shading(
                    p,
                    _style_str(
                        styles,
                        "ascii_background",
                        "asciiBackground",
                        default=_style_str(styles, "code_background", "codeBackground", default="#0f172a"),
                    ),
                )
                _style_paragraph_runs(
                    p,
                    font_name=_font_primary(
                        _style_str(
                            styles,
                            "ascii_font_family",
                            "asciiFontFamily",
                            default=code_font_name,
                        )
                    ),
                    size_pt=max(8.0, code_font_size - 1.0),
                    color_hex=_style_str(
                        styles,
                        "ascii_text",
                        "asciiText",
                        default=_style_str(styles, "code_text", "codeText", default="#1f2937"),
                    ),
                )
            elif node.type in {"image", "figure"}:
                stream = _resolve_image_stream(node.source, warnings)
                if stream:
                    p = document.add_paragraph()
                    p.alignment = _docx_alignment(node.align or "center")
                    run = p.add_run()
                    width_mm = max(20.0, min(170.0, 170.0 * ((node.scale or 100.0) / 100.0)))
                    try:
                        run.add_picture(stream, width=Mm(width_mm))
                    except Exception as exc:
                        warnings.append(f"Failed to render image '{node.source}': {exc}")
                        fallback = document.add_paragraph(f"[Image: {node.source}]")
                        fallback.alignment = _docx_alignment(node.align or "center")
                else:
                    fallback = document.add_paragraph(f"[Image: {node.source or 'missing source'}]")
                    fallback.alignment = _docx_alignment(node.align or "center")
                    _style_paragraph_runs(
                        fallback,
                        font_name=font_name,
                        size_pt=float(theme.bodyStyle.size or 12),
                        color_hex=body_color,
                    )

                caption_text = (node.caption or node.text or "").strip()
                if node.type == "figure" or caption_text:
                    number = _next_caption_number(caption_state, "figure")
                    cp = document.add_paragraph(f"Figure {number}: {caption_text or node.source or 'Image'}")
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cp.paragraph_format.space_before = Pt(2)
                    cp.paragraph_format.space_after = Pt(max(4.0, paragraph_after_pt))
                    cp.paragraph_format.line_spacing = line_spacing
                    _style_paragraph_runs(
                        cp,
                        font_name=font_name,
                        size_pt=max(11.0, float(theme.bodyStyle.size or 12) - 1.0),
                        color_hex="#475569",
                        italic=True,
                    )
            elif node.type == "table_caption":
                number = _next_caption_number(caption_state, "table")
                cp = document.add_paragraph(f"Table {number}: {node.text}")
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after = Pt(max(4.0, paragraph_after_pt))
                cp.paragraph_format.line_spacing = line_spacing
                _style_paragraph_runs(
                    cp,
                    font_name=font_name,
                    size_pt=max(11.0, float(theme.bodyStyle.size or 12) - 1.0),
                    color_hex="#475569",
                    italic=True,
                )
            elif node.type == "figure_caption":
                number = _next_caption_number(caption_state, "figure")
                cp = document.add_paragraph(f"Figure {number}: {node.text}")
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after = Pt(max(4.0, paragraph_after_pt))
                cp.paragraph_format.line_spacing = line_spacing
                _style_paragraph_runs(
                    cp,
                    font_name=font_name,
                    size_pt=max(11.0, float(theme.bodyStyle.size or 12) - 1.0),
                    color_hex="#475569",
                    italic=True,
                )
            elif node.type == "table":
                rows = node.rows or []
                if not rows:
                    continue
                cols = len(rows[0])
                table = document.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        value = row[c_idx] if c_idx < len(row) else ""
                        cell = table.cell(r_idx, c_idx)
                        cell.text = value
                        _set_cell_borders(
                            cell,
                            color=table_border_color,
                            width_pt=table_border_width,
                            style=table_border_style,
                        )
                        if r_idx == 0:
                            _set_cell_shading(cell, table_header_fill)
                        elif r_idx % 2 == 1:
                            _set_cell_shading(cell, table_odd_fill)
                        else:
                            _set_cell_shading(cell, table_even_fill)
                        for para in cell.paragraphs:
                            para.alignment = _docx_alignment(table_text_alignment)
                            para.paragraph_format.line_spacing = line_spacing
                            _style_paragraph_runs(
                                para,
                                font_name=font_name,
                                size_pt=float(theme.bodyStyle.size or 12),
                                color_hex=table_header_text if r_idx == 0 else body_color,
                                bold=True if r_idx == 0 else None,
                            )
            elif node.type == "separator":
                p = document.add_paragraph("")
                _apply_paragraph_separator(
                    p,
                    "top",
                    _style_str(styles, "table_border", "tableBorder", default="#D1D5DB"),
                    0.75,
                    "single",
                )
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            elif node.type == "pagebreak":
                document.add_page_break()

            if node.type != "pagebreak":
                has_written_content = True

        document.save(str(path))

        if security.removeMetadata:
            remove_docx_metadata(path)
        if security.disableEditingDocx:
            disable_docx_editing(path)

        return file_id, path, warnings

    def create_export_file(
        self,
        target_format: str,
        nodes: Sequence[Node],
        theme: ThemePayload,
        formatting: FormattingOptions,
        security: GenerateSecurityPayload,
    ) -> ExportResult:
        requested = target_format.lower()
        self.store.cleanup()

        if requested == "html":
            file_id, html_path = self.store.reserve_path("html")
            html = self.create_preview_html(nodes, theme, formatting, security)
            html_path.write_text(html, encoding="utf-8")
            return ExportResult(
                file_id=file_id,
                output_path=html_path,
                warnings=[],
                requested_format="html",
                actual_format="html",
                conversion_engine="native_html",
            )

        if requested == "md":
            file_id, md_path = self.store.reserve_path("md")
            md_path.write_text(to_markdown(nodes), encoding="utf-8")
            return ExportResult(
                file_id=file_id,
                output_path=md_path,
                warnings=[],
                requested_format="md",
                actual_format="md",
                conversion_engine="native_markdown",
            )

        if requested == "txt":
            file_id, txt_path = self.store.reserve_path("txt")
            txt_path.write_text(to_plain_text(nodes), encoding="utf-8")
            return ExportResult(
                file_id=file_id,
                output_path=txt_path,
                warnings=[],
                requested_format="txt",
                actual_format="txt",
                conversion_engine="native_text",
            )

        docx_id, docx_path, warnings = self.create_docx(nodes, theme, formatting, security)

        if requested == "docx":
            return ExportResult(
                file_id=docx_id,
                output_path=docx_path,
                warnings=warnings,
                requested_format="docx",
                actual_format="docx",
                conversion_engine="python_docx",
            )

        if requested == "pdf":
            file_id, pdf_path = self.store.reserve_path("pdf")
            attempts: list[str] = []
            external_methods = {"ilovepdf_api", "ilovepdf_automation"}
            ok, method = _convert_docx_to_pdf(docx_path, pdf_path)
            attempts.append(method)

            if not ok:
                ok, method = _convert_docx_to_pdf_ilovepdf_api(docx_path, pdf_path)
                attempts.append(method)

            if not ok:
                ok, method = _convert_docx_to_pdf_ilovepdf_automation(docx_path, pdf_path)
                attempts.append(method)

            if not ok:
                if _allow_low_fidelity_pdf_fallback():
                    html_preview = self.create_preview_html(
                        nodes,
                        theme,
                        formatting,
                        security,
                        include_running_blocks=False,
                    )
                    ok, method = _convert_html_to_pdf_weasyprint(
                        html_preview,
                        pdf_path,
                    )
                    attempts.append(method)

                if not ok:
                    ok, method = _convert_nodes_to_pdf_reportlab(
                        nodes,
                        theme,
                        formatting,
                        security,
                        pdf_path,
                    )
                    attempts.append(method)

            if not ok:
                ok, method = _write_emergency_pdf(
                    pdf_path,
                    "all converter backends unavailable",
                )
                attempts.append(method)
            if ok:
                warnings.extend(secure_pdf(pdf_path, security.passwordProtectPdf, security.removeMetadata))
                if method in {"weasyprint", "reportlab", "emergency_pdf"}:
                    warnings.append(
                        f"PDF generated via fallback renderer ({method}); install LibreOffice for closer DOCX-to-PDF fidelity."
                    )
                if method in external_methods:
                    warnings.append(f"PDF generated via iLovePDF external fallback ({method}).")
                return ExportResult(
                    file_id=file_id,
                    output_path=pdf_path,
                    warnings=warnings,
                    requested_format="pdf",
                    actual_format="pdf",
                    conversion_engine=method,
                    external_fallback_used=method in external_methods,
                )

            if pdf_path.exists():
                try:
                    pdf_path.unlink()
                except OSError:
                    pass

            raise RuntimeError(
                "PDF generation failed after all fallback strategies. "
                f"Converter detail: {' | '.join(attempts)}."
            )

        file_id, txt_path = self.store.reserve_path("txt")
        txt_path.write_text(to_plain_text(nodes), encoding="utf-8")
        return ExportResult(
            file_id=file_id,
            output_path=txt_path,
            warnings=[],
            requested_format=requested,
            actual_format="txt",
            conversion_engine="native_text_fallback",
        )
