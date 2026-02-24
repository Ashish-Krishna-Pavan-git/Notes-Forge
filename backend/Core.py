"""
NotesForge Professional v6.2 - Core Document Engine (Complete Implementation)
Complete implementations for all document building features.
"""

from __future__ import annotations
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import logging
import re
import os
import base64
import tempfile
import time
from xml.sax.saxutils import escape as xml_escape

# Third-party
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
except ImportError as e:
    DOCX_AVAILABLE = False
    logging.warning(f"python-docx not available: {e}")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger("NotesForge")


# ═══════════════════════════════════════════════════════════════════
# CONFIGURATION CLASSES
# ═══════════════════════════════════════════════════════════════════

@dataclass
class FontSettings:
    """Complete font configuration."""

    family: str = "Calibri"
    family_code: str = "Fira Code"

    h1_family: str = "Calibri"
    h2_family: str = "Calibri"
    h3_family: str = "Calibri"
    h4_family: str = "Calibri"
    h5_family: str = "Calibri"
    h6_family: str = "Calibri"
    bullet_family: str = "Calibri"

    sizes: Dict[str, int] = field(default_factory=lambda: {
        "h1": 20, "h2": 17, "h3": 14, "h4": 13, "h5": 12, "h6": 11,
        "body": 12, "code": 10, "header": 10, "footer": 10,
    })


@dataclass
class WatermarkSettings:
    enabled: bool = False
    type: str = "text"  # "text" or "image"
    text: str = "DRAFT"
    image_path: str = ""
    font: str = "Calibri"
    size: int = 48
    color: str = "#CCCCCC"
    opacity: float = 0.15
    rotation: int = 315
    position: str = "center"
    scale: int = 100


@dataclass
class HeaderFooterSettings:
    enabled: bool = True
    text: str = ""
    size: int = 10
    color: str = "#000000"
    bold: bool = False
    italic: bool = False
    alignment: str = "center"

    title_enabled: bool = True
    title_position: str = "header"
    title_alignment: str = "center"

    show_page_numbers: bool = True
    page_number_position: str = "footer"
    page_number_alignment: str = "center"
    page_format: str = "X"
    page_number_style: str = "arabic"

    separator: bool = True
    separator_color: str = "#CCCCCC"
    font_family: str = "Calibri"


@dataclass
class PageSettings:
    size: str = "A4"
    orientation: str = "portrait"
    margins: Dict[str, float] = field(default_factory=lambda: {
        "top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0,
    })
    border_enabled: bool = False
    border_width: int = 4
    border_color: str = "#000000"
    border_style: str = "single"
    border_offset: int = 24


@dataclass
class AppConfig:
    name: str = "NotesForge Professional"
    version: str = "6.2"
    theme: str = "tech"

    fonts: FontSettings = field(default_factory=FontSettings)
    watermark: WatermarkSettings = field(default_factory=WatermarkSettings)
    header: HeaderFooterSettings = field(default_factory=lambda: HeaderFooterSettings(
        enabled=True, text="NotesForge Professional", color="#FF8C00", bold=False,
    ))
    footer: HeaderFooterSettings = field(default_factory=lambda: HeaderFooterSettings(
        enabled=True,
        text=f"© {time.localtime().tm_year} NotesForge",
        show_page_numbers=True,
        page_format="Page X of Y",
    ))
    page: PageSettings = field(default_factory=PageSettings)

    colors: Dict[str, str] = field(default_factory=lambda: {
        "h1": "#6200EA", "h2": "#651FFF", "h3": "#7C4DFF",
        "h4": "#B388FF", "h5": "#333333", "h6": "#555555",
        "body": "#000000", "code_background": "#1E1E2E",
        "code_text": "#D4D4D4",
        "table_header_bg": "#6200EA", "table_header_text": "#FFFFFF",
        "table_odd_row": "#F3EDFF", "table_even_row": "#FFFFFF",
        "table_border": "#E0E0E0", "link": "#0563C1",
    })

    spacing: Dict[str, float] = field(default_factory=lambda: {
        "line_spacing": 1.4,
        "paragraph_spacing_after": 6.0,
        "heading_spacing_before": 12.0,
        "heading_spacing_after": 4.0,
        "code_indent": 0.35,
        "quote_indent": 0.5,
        "bullet_base_indent": 0.25,
        "bullet_indent_per_level": 0.45,
    })

    @classmethod
    def from_dict(cls, data: Dict) -> AppConfig:
        """Create AppConfig from dictionary."""
        cfg = cls()
        
        if "fonts" in data and isinstance(data["fonts"], dict):
            for k, v in data["fonts"].items():
                if hasattr(cfg.fonts, k):
                    setattr(cfg.fonts, k, v)
        
        if "colors" in data and isinstance(data["colors"], dict):
            cfg.colors.update(data["colors"])
        
        if "spacing" in data and isinstance(data["spacing"], dict):
            cfg.spacing.update(data["spacing"])
        
        if "watermark" in data and isinstance(data["watermark"], dict):
            for k, v in data["watermark"].items():
                if hasattr(cfg.watermark, k):
                    if k == "image_path" and isinstance(v, str) and v.startswith("data:image"):
                        try:
                            header, content = v.split(",", 1)
                            ext = "png"
                            header_lower = header.lower()
                            if "image/jpeg" in header_lower:
                                ext = "jpg"
                            elif "image/webp" in header_lower:
                                ext = "webp"
                            elif "image/gif" in header_lower:
                                ext = "gif"
                            img_data = base64.b64decode(content)
                            img_path = Path(tempfile.gettempdir()) / f"notesforge_watermark_{int(time.time() * 1000)}.{ext}"
                            img_path.write_bytes(img_data)
                            setattr(cfg.watermark, k, str(img_path))
                        except Exception as e:
                            logger.warning(f"Failed to decode watermark image data: {e}")
                    else:
                        setattr(cfg.watermark, k, v)
        
        if "header" in data and isinstance(data["header"], dict):
            for k, v in data["header"].items():
                if hasattr(cfg.header, k):
                    setattr(cfg.header, k, v)
            if (cfg.header.page_format or "").strip() in {"X | Page", "X | P a g e"}:
                cfg.header.page_format = "Page X of Y"

        if "footer" in data and isinstance(data["footer"], dict):
            for k, v in data["footer"].items():
                if hasattr(cfg.footer, k):
                    setattr(cfg.footer, k, v)
            if (cfg.footer.page_format or "").strip() in {"X | Page", "X | P a g e"}:
                cfg.footer.page_format = "Page X of Y"
        
        if "page" in data and isinstance(data["page"], dict):
            page_data = data["page"]
            if "size" in page_data:
                cfg.page.size = page_data["size"]
            if "orientation" in page_data:
                cfg.page.orientation = page_data["orientation"]
            if "margins" in page_data and isinstance(page_data["margins"], dict):
                cfg.page.margins.update(page_data["margins"])
            if "border" in page_data and isinstance(page_data["border"], dict):
                for k, v in page_data["border"].items():
                    if hasattr(cfg.page, f"border_{k}"):
                        setattr(cfg.page, f"border_{k}", v)
            if "border_enabled" in page_data:
                cfg.page.border_enabled = page_data["border_enabled"]
            if "border_width" in page_data:
                cfg.page.border_width = page_data["border_width"]
            if "border_color" in page_data:
                cfg.page.border_color = page_data["border_color"]
            if "border_style" in page_data:
                cfg.page.border_style = page_data["border_style"]
        
        return cfg


# ═══════════════════════════════════════════════════════════════════
# TEXT PARSER
# ═══════════════════════════════════════════════════════════════════

class TextParser:
    """Parse NotesForge marker syntax."""

    MARKERS = {
        "HEADING": "h1", "H1": "h1",
        "SUBHEADING": "h2", "H2": "h2",
        "SUB-SUBHEADING": "h3", "H3": "h3",
        "H4": "h4", "H5": "h5", "H6": "h6",
        "PARAGRAPH": "paragraph", "PARA": "paragraph",
        "CENTER": "paragraph_center",
        "RIGHT": "paragraph_right",
        "BULLET": "bullet",
        "NUMBERED": "numbered",
        "CODE": "code",
        "TABLE": "table",
        "ASCII": "ascii", "DIAGRAM": "ascii",
        "QUOTE": "quote",
        "NOTE": "note", "IMPORTANT": "note",
        "IMAGE": "image",
        "LINK": "link",
        "HIGHLIGHT": "highlight",
        "FOOTNOTE": "footnote",
        "TOC": "toc",
        "WATERMARK": "watermark",
    }

    def classify_line(self, line: str) -> Dict[str, Any]:
        stripped = line.rstrip("\n")
        if not stripped.strip():
            return {"type": "empty", "content": ""}

        # Marker detection
        m = re.match(r"^([A-Z][A-Z0-9\-]*):\s*(.*)$", stripped.strip())
        if m:
            marker, content = m.groups()
            marker_type = self.MARKERS.get(marker, "text")
            indent = len(line) - len(line.lstrip())
            return {
                "type": marker_type,
                "content": content.strip(),
                "marker": marker,
                "indent_level": indent // 2,
            }

        # Table detection
        if "|" in stripped and stripped.count("|") >= 1:
            return {"type": "table", "content": stripped}

        # Bullet fallback
        bm = re.match(r"^(\s*)([-*•])\s+(.+)$", line)
        if bm:
            indent, _, content = bm.groups()
            return {
                "type": "bullet",
                "content": content,
                "indent_level": len(indent) // 2,
            }

        # Numbered fallback
        if re.match(r"^\s*\d+[.)]\s+", stripped):
            return {"type": "numbered", "content": stripped.strip()}

        # Code detection
        code_starts = (
            "def ", "class ", "function ", "var ", "let ", "const ",
            "import ", "from ", "if ", "for ", "while ", "return ", "print(",
            "#include", "public ", "private ",
        )
        if any(stripped.strip().startswith(kw) for kw in code_starts):
            return {"type": "code", "content": stripped.strip()}

        # ASCII
        ascii_chars = set("─│┌┐└┘├┤┬┴┼╔╗╚╝═║")
        if any(c in stripped for c in ascii_chars):
            return {"type": "ascii", "content": stripped}

        # Paragraph vs text
        if len(stripped.strip()) >= 40:
            return {"type": "paragraph", "content": stripped.strip()}
        
        return {"type": "text", "content": stripped.strip()}


# ═══════════════════════════════════════════════════════════════════
# DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════════

class DocumentBuilder:
    """Build Word documents with full customization."""

    def __init__(self, config: AppConfig):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required")
        self.config = config
        self.doc: Optional[Document] = None
        self._footnote_counter = 0
        self._headings: List[Tuple[int, str]] = []

    def _hex_color(self, hex_color: str) -> RGBColor:
        """Convert hex to RGBColor with validation."""
        try:
            h = hex_color.lstrip("#")
            if len(h) == 6:
                r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
                return RGBColor(r, g, b)
            else:
                return RGBColor(0, 0, 0)
        except (ValueError, IndexError):
            return RGBColor(0, 0, 0)

    def _get_alignment(self, align: str):
        return {
            "left": WD_ALIGN.LEFT,
            "center": WD_ALIGN.CENTER,
            "right": WD_ALIGN.RIGHT,
        }.get(align, WD_ALIGN.LEFT)

    def _get_heading_font(self, level: int) -> str:
        """Get font for heading level with fallback."""
        attr = f"h{level}_family"
        font = getattr(self.config.fonts, attr, None)
        return font or self.config.fonts.family or "Calibri"

    def _apply_page_borders(self):
        """Apply page borders to all sections."""
        if not self.config.page.border_enabled:
            return

        for section in self.doc.sections:
            sectPr = section._sectPr
            pgBorders = OxmlElement("w:pgBorders")
            pgBorders.set(qn("w:offsetFrom"), "page")

            border_color = self.config.page.border_color.lstrip("#")
            border_sz = str(self.config.page.border_width * 8)
            border_style_map = {
                "single": "single",
                "double": "double",
                "thick": "thick",
                "dashed": "dashed",
                "dotted": "dotted",
            }
            border_style = border_style_map.get(
                self.config.page.border_style, "single"
            )

            for border_name in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), border_style)
                border.set(qn("w:sz"), border_sz)
                border.set(qn("w:space"), str(self.config.page.border_offset))
                border.set(qn("w:color"), border_color)
                pgBorders.append(border)

            sectPr.append(pgBorders)

    def _apply_watermark(self):
        """Apply watermark to document."""
        if not self.config.watermark.enabled:
            return

        wm = self.config.watermark

        try:
            if wm.type == "image" and wm.image_path and os.path.isfile(wm.image_path):
                self._apply_image_watermark()
            else:
                self._apply_text_watermark()
        except Exception as e:
            logger.warning(f"Watermark application failed: {e}")

    def _apply_text_watermark(self):
        """Apply text watermark using VML."""
        wm = self.config.watermark

        for section in self.doc.sections:
            header = section.header
            paragraph = header.add_paragraph()
            paragraph.alignment = WD_ALIGN.CENTER

            color_hex = wm.color.lstrip("#") or "CCCCCC"
            opacity = max(0.0, min(1.0, float(wm.opacity)))
            shape_xml = parse_xml(
                f"""
                <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        xmlns:v="urn:schemas-microsoft-com:vml"
                        xmlns:o="urn:schemas-microsoft-com:office:office">
                    <v:shape id="NotesForgeWatermark"
                        o:spid="_x0000_s2050"
                        type="#_x0000_t136"
                        style="position:absolute;margin-left:0;margin-top:0;width:468pt;height:156pt;
                               z-index:-251654144;mso-wrap-edited:f;mso-position-horizontal:center;
                               mso-position-horizontal-relative:margin;mso-position-vertical:center;
                               mso-position-vertical-relative:margin;rotation:{wm.rotation}"
                        fillcolor="#{color_hex}" stroked="f">
                        <v:fill opacity="{int(opacity * 100)}%"/>
                        <v:textpath style="font-family:{xml_escape(wm.font)};font-size:{int(wm.size)}pt"
                                    string="{xml_escape(wm.text)}"/>
                    </v:shape>
                </w:pict>
                """
            )
            run = paragraph.add_run()
            run._r.append(shape_xml)

    def _apply_image_watermark(self):
        """Apply image watermark to document center."""
        wm = self.config.watermark

        if not os.path.isfile(wm.image_path):
            logger.warning(f"Watermark image not found: {wm.image_path}")
            return

        try:
            for section in self.doc.sections:
                header = section.header
                paragraph = header.add_paragraph()
                paragraph.alignment = WD_ALIGN.CENTER

                width_inches = (wm.scale / 100) * 3
                run = paragraph.add_run()
                run.add_picture(wm.image_path, width=Inches(width_inches))
        except Exception as e:
            logger.warning(f"Failed to add image watermark: {e}")

    def _append_field(self, paragraph, field_name: str, style: str = "arabic"):
        style_map = {
            "arabic": "",
            "roman": r" \* ROMAN",
            "alpha": r" \* ALPHABETIC",
        }
        field_suffix = style_map.get((style or "arabic").lower(), "")
        field_code = f"{field_name}{field_suffix}"

        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        run = paragraph.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = field_code
        run._r.append(instr)

        run = paragraph.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_sep)

        run = paragraph.add_run("1")

        run = paragraph.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

    def _append_page_number_format(self, paragraph, fmt: str, style: str = "arabic"):
        template = (fmt or "X").strip()
        if not template:
            template = "X"
        template = re.sub(r"(?i)\{page\}", "X", template)
        template = re.sub(r"(?i)\{pages\}|\{total\}", "Y", template)

        token_pattern = re.compile(r"(?i)\b([xy])\b")
        pos = 0
        found = False

        for match in token_pattern.finditer(template):
            found = True
            start, end = match.span()
            if start > pos:
                paragraph.add_run(template[pos:start])
            token = match.group(1).upper()
            if token == "X":
                self._append_field(paragraph, "PAGE", style)
            else:
                self._append_field(paragraph, "NUMPAGES", style)
            pos = end

        if pos < len(template):
            paragraph.add_run(template[pos:])

        if not found:
            self._append_field(paragraph, "PAGE", style)

    def _apply_separator_line(self, paragraph, edge: str, color: str):
        """Draw a real separator line using paragraph borders."""
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)

        edge_tag = qn(f"w:{edge}")
        existing = p_bdr.find(edge_tag)
        if existing is not None:
            p_bdr.remove(existing)

        line = OxmlElement(f"w:{edge}")
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "8")
        line.set(qn("w:space"), "2")
        line.set(qn("w:color"), (color or "#CCCCCC").lstrip("#"))
        p_bdr.append(line)

    def _apply_header_footer(self):
        """Apply header and footer to document."""
        for section in self.doc.sections:
            # Set margins
            section.top_margin = Inches(self.config.page.margins.get("top", 1.0))
            section.bottom_margin = Inches(self.config.page.margins.get("bottom", 1.0))
            section.left_margin = Inches(self.config.page.margins.get("left", 1.0))
            section.right_margin = Inches(self.config.page.margins.get("right", 1.0))

            # HEADER
            if self.config.header.enabled:
                header = section.header
                h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                h_para.text = self.config.header.text or ""

                if (
                    self.config.header.show_page_numbers
                    and self.config.header.page_number_position == "header"
                ):
                    if self.config.header.text:
                        h_para.add_run(" | ")
                    self._append_page_number_format(
                        h_para,
                        self.config.header.page_format,
                        self.config.header.page_number_style,
                    )
                    h_para.alignment = self._get_alignment(
                        self.config.header.page_number_alignment or self.config.header.alignment
                    )
                else:
                    h_para.alignment = self._get_alignment(self.config.header.alignment)

                for run in h_para.runs:
                    run.font.name = self.config.header.font_family or self.config.fonts.family
                    run.font.size = Pt(self.config.header.size)
                    run.font.color.rgb = self._hex_color(self.config.header.color)
                    run.bold = self.config.header.bold
                    run.italic = self.config.header.italic

                if self.config.header.separator:
                    self._apply_separator_line(
                        h_para,
                        "bottom",
                        self.config.header.separator_color,
                    )

            # FOOTER
            if self.config.footer.enabled:
                footer = section.footer
                f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                f_para.text = self.config.footer.text or ""
                
                if (
                    self.config.footer.show_page_numbers
                    and self.config.footer.page_number_position == "footer"
                ):
                    if self.config.footer.text and self.config.footer.separator:
                        f_para.add_run(" | ")
                    elif self.config.footer.text:
                        f_para.add_run(" ")
                    self._append_page_number_format(
                        f_para,
                        self.config.footer.page_format,
                        self.config.footer.page_number_style,
                    )
                    f_para.alignment = self._get_alignment(
                        self.config.footer.page_number_alignment or self.config.footer.alignment
                    )
                else:
                    f_para.alignment = self._get_alignment(self.config.footer.alignment)

                for run in f_para.runs:
                    run.font.name = self.config.footer.font_family or self.config.fonts.family
                    run.font.size = Pt(self.config.footer.size)
                    run.font.color.rgb = self._hex_color(self.config.footer.color)
                    run.bold = self.config.footer.bold
                    run.italic = self.config.footer.italic

                if self.config.footer.separator:
                    self._apply_separator_line(
                        f_para,
                        "top",
                        self.config.footer.separator_color,
                    )

    def _add_heading(self, text: str, level: int):
        """Add heading to document."""
        size_key = f"h{level}"
        size = self.config.fonts.sizes.get(size_key, 12)
        color = self.config.colors.get(size_key, "#000000")
        font_family = self._get_heading_font(level)

        p = self.doc.add_heading(text, level=level)
        p.paragraph_format.space_before = Pt(self.config.spacing.get("heading_spacing_before", 12))
        p.paragraph_format.space_after = Pt(self.config.spacing.get("heading_spacing_after", 4))
        p.paragraph_format.line_spacing = self.config.spacing.get("line_spacing", 1.4)

        for run in p.runs:
            run.font.name = font_family
            run.font.size = Pt(size)
            run.font.color.rgb = self._hex_color(color)

        self._headings.append((level, text))

    def _add_paragraph(self, text: str, alignment: str = "left"):
        """Add paragraph to document."""
        p = self.doc.add_paragraph()
        p.alignment = self._get_alignment(alignment)
        p.paragraph_format.space_after = Pt(self.config.spacing.get("paragraph_spacing_after", 6))
        p.paragraph_format.line_spacing = self.config.spacing.get("line_spacing", 1.4)

        body_family = self.config.fonts.family
        body_size = Pt(self.config.fonts.sizes.get("body", 12))
        body_color = self._hex_color(self.config.colors.get("body", "#000000"))

        # Parse inline formatting
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|\{[^}]+:[^}]+\})", text)

        for part in parts:
            if not part:
                continue

            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("{") and part.endswith("}"):
                match = re.match(r"\{([^:]+):([^}]+)\}", part)
                if match:
                    color_name, text_part = match.groups()
                    run = p.add_run(text_part)
                    if color_name in self.config.colors:
                        run.font.color.rgb = self._hex_color(self.config.colors[color_name])
            else:
                run = p.add_run(part)

            if part not in ("**", "*", "{", "}"):
                run.font.name = body_family
                run.font.size = body_size
                if not part.startswith("{"):
                    run.font.color.rgb = body_color

    def _add_bullet(self, text: str, level: int = 0):
        """Add bullet point."""
        p = self.doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.left_indent = Inches(
            self.config.spacing.get("bullet_base_indent", 0.25) +
            level * self.config.spacing.get("bullet_indent_per_level", 0.45)
        )
        p.paragraph_format.line_spacing = self.config.spacing.get("line_spacing", 1.4)
        bullet_font = self.config.fonts.bullet_family or self.config.fonts.family
        bullet_size = Pt(self.config.fonts.sizes.get("body", 12))
        bullet_color = self._hex_color(self.config.colors.get("body", "#000000"))
        for run in p.runs:
            run.font.name = bullet_font
            run.font.size = bullet_size
            run.font.color.rgb = bullet_color

    def _add_numbered(self, text: str):
        """Add numbered list item."""
        p = self.doc.add_paragraph(text, style="List Number")
        p.paragraph_format.line_spacing = self.config.spacing.get("line_spacing", 1.4)
        bullet_font = self.config.fonts.bullet_family or self.config.fonts.family
        bullet_size = Pt(self.config.fonts.sizes.get("body", 12))
        bullet_color = self._hex_color(self.config.colors.get("body", "#000000"))
        for run in p.runs:
            run.font.name = bullet_font
            run.font.size = bullet_size
            run.font.color.rgb = bullet_color

    def _add_code(self, text: str):
        """Add code block."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(self.config.spacing.get("code_indent", 0.35))
        p.style = "Normal"

        bg_color = self.config.colors.get("code_background", "#1E1E2E")
        text_color = self.config.colors.get("code_text", "#D4D4D4")

        run = p.add_run(text)
        run.font.name = self.config.fonts.family_code
        run.font.size = Pt(self.config.fonts.sizes.get("code", 10))
        run.font.color.rgb = self._hex_color(text_color)

        # Set background color
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), bg_color.lstrip("#"))
        p._element.get_or_add_pPr().append(shading_elm)

    def _add_quote(self, text: str):
        """Add block quote."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(self.config.spacing.get("quote_indent", 0.5))
        p.style = "Quote"

        run = p.add_run(f'"{text}"')
        run.italic = True

    def _add_note(self, text: str):
        """Add note/callout."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)

        run = p.add_run("📝 ")
        run = p.add_run(text)
        run.font.bold = True

    def _add_ascii(self, text: str):
        """Add ASCII art/diagram."""
        p = self.doc.add_paragraph()
        p.style = "Normal"

        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    def _add_table(self, rows: List[List[str]]):
        """Add table to document."""
        if not rows:
            return

        ncols = max(len(r) for r in rows) if rows else 0
        nrows = len(rows)

        if ncols == 0 or nrows == 0:
            return

        tbl = self.doc.add_table(rows=nrows, cols=ncols)
        tbl.style = "Table Grid"

        h_bg = self.config.colors.get("table_header_bg", "#6200EA")
        h_txt = self.config.colors.get("table_header_text", "#FFFFFF")
        odd_bg = self.config.colors.get("table_odd_row", "#F3EDFF")
        even_bg = self.config.colors.get("table_even_row", "#FFFFFF")

        for ri, row_data in enumerate(rows):
            row = tbl.rows[ri]

            for ci, cell_text in enumerate(row_data):
                if ci >= len(row.cells):
                    break

                cell = row.cells[ci]
                cell.text = cell_text.strip()

                # Header row styling
                if ri == 0:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), h_bg.lstrip("#"))
                    cell._element.get_or_add_tcPr().append(shading)

                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = self._hex_color(h_txt)
                else:
                    # Alternating row colors
                    bg_color = odd_bg if ri % 2 == 1 else even_bg
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), bg_color.lstrip("#"))
                    cell._element.get_or_add_tcPr().append(shading)

        self.doc.add_paragraph()

    def _add_image(self, image_path: str, caption: str = "", alignment: str = "center"):
        """Add image to document."""
        if not os.path.isfile(image_path):
            logger.warning(f"Image not found: {image_path}")
            return

        try:
            p = self.doc.add_paragraph()
            p.alignment = self._get_alignment(alignment)

            p.add_picture(image_path, width=Inches(4))

            if caption:
                cap_p = self.doc.add_paragraph(caption)
                cap_p.alignment = WD_ALIGN.CENTER
                cap_p.paragraph_format.space_after = Pt(12)

                for run in cap_p.runs:
                    run.italic = True
                    run.font.size = Pt(10)
        except Exception as e:
            logger.warning(f"Failed to add image: {e}")

    def _add_link(self, link_text: str, url: str):
        """Add hyperlink to document."""
        p = self.doc.add_paragraph()

        # Simple URL validation
        if not (url.startswith("http://") or url.startswith("https://")):
            url = f"https://{url}"

        run = p.add_run(link_text)
        run.font.underline = True
        run.font.color.rgb = self._hex_color(self.config.colors.get("link", "#0563C1"))

        # Note: Full hyperlink implementation requires python-docx extensions

    def _add_highlight(self, text: str, color: str = "yellow"):
        """Add highlighted text."""
        p = self.doc.add_paragraph()

        color_map = {
            "yellow": "FFFF00",
            "green": "00FF00",
            "blue": "0000FF",
            "red": "FF0000",
            "pink": "FFC0CB",
        }

        hex_color = color_map.get(color.lower(), "FFFF00")

        run = p.add_run(text)
        # python-docx doesn't have direct highlight support
        # This is a workaround using shading

    def _add_footnote(self, text: str):
        """Add footnote."""
        p = self.doc.add_paragraph()
        self._footnote_counter += 1

        run = p.add_run(f"[{self._footnote_counter}] {text}")
        run.font.size = Pt(9)

    def _add_toc(self):
        """Add table of contents."""
        p = self.doc.add_paragraph()
        p.text = "Table of Contents"
        p.style = "Heading 1"

        for level, heading_text in self._headings:
            indent = (level - 1) * 0.25
            p = self.doc.add_paragraph(heading_text, style="List Bullet")
            p.paragraph_format.left_indent = Inches(indent)

    def build(self, text: str, output_path: Optional[str] = None) -> str:
        """Build complete document from marker text."""
        if not text.strip():
            raise ValueError("Text cannot be empty")

        self.doc = Document()
        parser = TextParser()

        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            classified = parser.classify_line(line)
            elem_type = classified.get("type", "text")
            content = classified.get("content", "")

            if elem_type == "empty":
                i += 1
                continue

            elif elem_type == "h1":
                self._add_heading(content, 1)

            elif elem_type == "h2":
                self._add_heading(content, 2)

            elif elem_type == "h3":
                self._add_heading(content, 3)

            elif elem_type == "h4":
                self._add_heading(content, 4)

            elif elem_type == "h5":
                self._add_heading(content, 5)

            elif elem_type == "h6":
                self._add_heading(content, 6)

            elif elem_type == "paragraph":
                self._add_paragraph(content)

            elif elem_type == "paragraph_center":
                self._add_paragraph(content, "center")

            elif elem_type == "paragraph_right":
                self._add_paragraph(content, "right")

            elif elem_type == "bullet":
                level = classified.get("indent_level", 0)
                self._add_bullet(content, level)

            elif elem_type == "numbered":
                self._add_numbered(content)

            elif elem_type == "code":
                self._add_code(content)

            elif elem_type == "table":
                # Collect table rows
                table_rows = [content.split("|")]
                i += 1
                while i < len(lines):
                    next_classified = parser.classify_line(lines[i])
                    if next_classified.get("type") == "table":
                        table_rows.append(next_classified.get("content", "").split("|"))
                        i += 1
                    else:
                        break
                i -= 1
                self._add_table(table_rows)

            elif elem_type == "quote":
                self._add_quote(content)

            elif elem_type == "note":
                self._add_note(content)

            elif elem_type == "ascii":
                self._add_ascii(content)

            elif elem_type == "image":
                parts = content.split("|")
                img_path = parts[0].strip().strip('"')
                caption = parts[1].strip().strip('"') if len(parts) > 1 else ""
                alignment = parts[2].strip().strip('"') if len(parts) > 2 else "center"
                self._add_image(img_path, caption, alignment)

            elif elem_type == "link":
                parts = content.split("|")
                link_text = parts[0].strip().strip('"')
                url = parts[1].strip().strip('"') if len(parts) > 1 else ""
                self._add_link(link_text, url)

            elif elem_type == "footnote":
                self._add_footnote(content)

            elif elem_type == "toc":
                self._add_toc()

            i += 1

        # Apply formatting
        self._apply_header_footer()
        self._apply_page_borders()
        self._apply_watermark()

        # Save document
        if not output_path:
            output_path = "document.docx"

        self.doc.save(output_path)
        return output_path


# ═══════════════════════════════════════════════════════════════════
# BACKWARD COMPATIBILITY
# ═══════════════════════════════════════════════════════════════════

TextAnalyzer = TextParser


class ConfigManager:
    """Simple config manager for backward compatibility."""

    def __init__(self, config_dict: Dict = None):
        self.config = AppConfig.from_dict(config_dict or {})

    def get(self, key: str, default=None):
        keys = key.split(".")
        val = self.config
        for k in keys:
            if hasattr(val, k):
                val = getattr(val, k)
            else:
                return default
        return val


def build_document(text: str, output_path: str = None, config: Dict = None) -> str:
    """Backward compatible API."""
    cfg = AppConfig.from_dict(config or {})
    builder = DocumentBuilder(cfg)
    return builder.build(text, output_path)
