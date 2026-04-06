import unittest
import time
from unittest.mock import patch
import io
import zipfile

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter

from app.main import create_app
from app.templates_repo import SAMPLE_EXAMPLE
from app.themes import PROFESSIONAL_THEME


class ApiIntegrationTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.client = TestClient(create_app())

    def _sample_pdf_bytes(self) -> bytes:
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        buffer = io.BytesIO()
        writer.write(buffer)
        return buffer.getvalue()

    def _sample_docx_bytes(self) -> bytes:
        buffer = io.BytesIO()
        document = Document()
        document.add_paragraph("NotesForge processing test")
        document.save(buffer)
        return buffer.getvalue()

    def test_health_endpoints(self) -> None:
        health = self.client.get("/api/health")
        self.assertEqual(health.status_code, 200)
        self.assertEqual(health.json(), {"status": "ok"})

        parser_health = self.client.get("/api/health/parser")
        self.assertEqual(parser_health.status_code, 200)
        self.assertEqual(parser_health.json()["parser"], "ok")

    def test_templates_contract(self) -> None:
        response = self.client.get("/api/templates")
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        ids = {item["id"] for item in payload}
        self.assertTrue({"assignment", "resume", "report", "meeting", "cybersec"}.issubset(ids))
        names = {item["name"] for item in payload}
        self.assertTrue(
            {
                "Project Report Template",
                "Research Paper Template",
                "Study Notes Template",
                "Technical Documentation Template",
                "Assignment Template",
            }.issubset(names)
        )

    def test_markers_catalog_contract(self) -> None:
        response = self.client.get("/api/markers")
        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body.get("success"))
        markers = body.get("markers", [])
        self.assertTrue(isinstance(markers, list) and len(markers) > 10)
        keys = {item.get("key") for item in markers}
        self.assertTrue(
            {
                "H1",
                "PARAGRAPH",
                "TABLE",
                "CAPTION",
                "ASCII",
                "FIGURE",
                "CHAPTER",
                "PAGEBREAK",
                "CHECKLIST",
                "EQUATION",
                "SEPARATOR",
            }.issubset(keys)
        )
        self.assertNotIn("WATERMARK", keys)
        ascii_item = next((item for item in markers if item.get("key") == "ASCII"), {})
        self.assertIn("DIAGRAM", ascii_item.get("aliases", []))
        paragraph_item = next((item for item in markers if item.get("key") == "PARAGRAPH"), {})
        self.assertIn("BODY", paragraph_item.get("aliases", []))

    def test_themes_include_required_v7_builtins(self) -> None:
        response = self.client.get("/api/themes")
        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body.get("success"))
        themes = body.get("themes", {})
        names = {v.get("name") for v in themes.values() if isinstance(v, dict)}
        self.assertTrue(
            {
                "Academic Classic",
                "University Blue",
                "Engineering Report",
                "Clean Research",
                "Modern Minimal",
                "Corporate White",
                "Dark Technical",
                "Elegant Thesis",
                "Lecture Notes",
                "Professional Docs",
            }.issubset(names)
        )

    def test_preflight_cors(self) -> None:
        for origin in (
            "http://localhost:5173",
            "http://127.0.0.1:5173",
            "https://notes-forge.onrender.com",
        ):
            response = self.client.options(
                "/api/templates",
                headers={
                    "Origin": origin,
                    "Access-Control-Request-Method": "GET",
                },
            )
            self.assertIn(response.status_code, (200, 204))
            self.assertEqual(response.headers.get("access-control-allow-origin"), origin)

    def test_preflight_cors_rejects_unknown_origin(self) -> None:
        response = self.client.options(
            "/api/templates",
            headers={
                "Origin": "https://malicious.example",
                "Access-Control-Request-Method": "GET",
            },
        )
        self.assertIn(response.status_code, (200, 204, 400))
        self.assertNotEqual(
            response.headers.get("access-control-allow-origin"),
            "https://malicious.example",
        )

    def test_security_headers_present(self) -> None:
        response = self.client.get("/api/health")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers.get("x-content-type-options"), "nosniff")
        self.assertEqual(response.headers.get("x-frame-options"), "DENY")
        self.assertEqual(response.headers.get("cache-control"), "no-store")

    def test_preview_contract(self) -> None:
        response = self.client.post(
            "/api/preview",
            json={
                "content": SAMPLE_EXAMPLE,
                "theme": PROFESSIONAL_THEME.model_dump(),
                "formattingOptions": {
                    "margins": {"top": 25, "bottom": 25, "left": 25, "right": 25},
                    "lineSpacing": 1.4,
                },
                "security": {
                    "removeMetadata": False,
                    "pageNumberMode": "page_x",
                },
            },
        )
        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertIn("previewHtml", body)
        self.assertIn("structure", body)
        self.assertGreaterEqual(body["structure"]["headingCount"], 1)

    def test_generate_and_download_docx(self) -> None:
        response = self.client.post(
            "/api/generate",
            json={
                "content": SAMPLE_EXAMPLE,
                "theme": PROFESSIONAL_THEME.model_dump(),
                "format": "docx",
                "filename": "notesforge_output",
                "security": {"disableEditingDocx": False, "removeMetadata": True},
                "templateId": "assignment",
            },
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload.get("success"))
        self.assertIn("downloadUrl", payload)
        self.assertIn("fileId", payload)
        self.assertEqual(payload.get("requestedFormat"), "docx")
        self.assertEqual(payload.get("actualFormat"), "docx")
        self.assertEqual(payload.get("conversionEngine"), "python_docx")
        self.assertFalse(payload.get("externalFallbackUsed"))
        self.assertEqual(payload.get("filename"), "notesforge_output.docx")

        download = self.client.get(payload["downloadUrl"])
        self.assertEqual(download.status_code, 200)
        self.assertTrue(
            download.headers["content-type"].startswith(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        )

    def test_generate_docx_reflects_theme_border_alignment_and_table_styles(self) -> None:
        custom_theme = PROFESSIONAL_THEME.model_dump()
        custom_theme["styles"] = {
            "page_border_enabled": True,
            "page_border_color": "#B91C1C",
            "page_border_width": 2,
            "page_border_style": "double",
            "header_alignment": "left",
            "footer_alignment": "right",
            "page_number_alignment": "right",
            "page_number_position": "footer",
            "footer_color": "#0055AA",
            "header_show_page_numbers": False,
            "footer_show_page_numbers": True,
            "table_header_text": "#FFFFFF",
            "table_odd_row": "#F3EDFF",
        }
        content = 'H1: "Theme Test"\nTABLE: "A | B"\nTABLE: "1 | 2"'
        response = self.client.post(
            "/api/generate",
            json={
                "content": content,
                "theme": custom_theme,
                "format": "docx",
                "filename": "theme_reflect",
                "security": {
                    "headerText": "Header Demo",
                    "footerText": "Footer Demo",
                    "pageNumberMode": "page_x_of_y",
                    "disableEditingDocx": False,
                    "removeMetadata": False,
                },
            },
        )
        self.assertEqual(response.status_code, 200)
        download = self.client.get(response.json()["downloadUrl"])
        self.assertEqual(download.status_code, 200)

        with zipfile.ZipFile(io.BytesIO(download.content)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", "ignore")
            header_xml = archive.read("word/header1.xml").decode("utf-8", "ignore")
            footer_xml = archive.read("word/footer1.xml").decode("utf-8", "ignore")

        document_xml_u = document_xml.upper()
        self.assertIn("W:PGBORDERS", document_xml_u)
        self.assertIn('w:jc w:val="left"', header_xml)
        self.assertIn('w:jc w:val="right"', footer_xml)
        self.assertIn("NUMPAGES", footer_xml)
        self.assertIn("PAGE", footer_xml)
        self.assertIn('w:color w:val="0055AA"', footer_xml)
        self.assertIn('W:FILL="F6F6F6"', document_xml_u)
        self.assertIn('W:FILL="F3EDFF"', document_xml_u)

    def test_generate_and_download_textual_formats(self) -> None:
        for fmt, ctype_prefix in [
            ("html", "text/html"),
            ("md", "text/markdown"),
            ("txt", "text/plain"),
        ]:
            response = self.client.post(
                "/api/generate",
                json={
                    "content": SAMPLE_EXAMPLE,
                    "theme": PROFESSIONAL_THEME.model_dump(),
                    "format": fmt,
                    "filename": f"notesforge_output_{fmt}",
                    "security": {"disableEditingDocx": False, "removeMetadata": False},
                    "templateId": "assignment",
                },
            )
            self.assertEqual(response.status_code, 200)
            payload = response.json()
            self.assertEqual(payload.get("requestedFormat"), fmt)
            self.assertEqual(payload.get("actualFormat"), fmt)
            self.assertTrue(payload.get("filename", "").endswith(f".{fmt}"))
            download = self.client.get(payload["downloadUrl"])
            self.assertEqual(download.status_code, 200)
            self.assertTrue(download.headers["content-type"].startswith(ctype_prefix))

    def test_generate_docx_marker_coverage_justify_ascii_code(self) -> None:
        content = (
            "H1: Main Title\n"
            "H2: Section\n"
            "H3: Subsection\n"
            "H4: Topic\n"
            "H5: Topic 5\n"
            "H6: Topic 6\n"
            "PARAGRAPH: left paragraph\n"
            "CENTER: centered paragraph\n"
            "RIGHT: right paragraph\n"
            "JUSTIFY: this paragraph should be justified\n"
            "with continuation text\n"
            "BULLET: bullet one\n"
            "- bullet two\n"
            "NUMBERED: first item\n"
            "2. second item\n"
            "CODE: print('hello')\n"
            "for i in range(2):\n"
            "    print(i)\n"
            "ASCII: +---+\n"
            "| A |\n"
            "+---+\n"
            "TABLE: Col1 | Col2\n"
            "TABLE: V1 | V2\n"
            "PAGEBREAK:\n"
            "PARAGRAPH: after break\n"
        )
        response = self.client.post(
            "/api/generate",
            json={
                "content": content,
                "theme": PROFESSIONAL_THEME.model_dump(),
                "format": "docx",
                "filename": "marker_coverage",
                "security": {"disableEditingDocx": False, "removeMetadata": False},
            },
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload.get("requestedFormat"), "docx")
        self.assertEqual(payload.get("actualFormat"), "docx")

        download = self.client.get(payload["downloadUrl"])
        self.assertEqual(download.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(download.content)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", "ignore")

        self.assertIn("this paragraph should be justified", document_xml)
        self.assertIn("with continuation text", document_xml)
        self.assertIn("print('hello')", document_xml)
        self.assertIn("+---+", document_xml)
        self.assertIn("| A |", document_xml)
        self.assertIn("Col1", document_xml)
        self.assertIn("V1", document_xml)
        self.assertIn("after break", document_xml)
        self.assertIn('w:jc w:val="both"', document_xml)

    def test_config_endpoints(self) -> None:
        current = self.client.get("/api/config")
        self.assertEqual(current.status_code, 200)
        body = current.json()
        self.assertTrue(body.get("success"))
        self.assertIn("config", body)

        update = self.client.post(
            "/api/config/update",
            json={"path": "spacing.line_spacing", "value": 1.6},
        )
        self.assertEqual(update.status_code, 200)
        self.assertEqual(update.json(), {"success": True})

        after = self.client.get("/api/config")
        self.assertEqual(after.status_code, 200)
        self.assertEqual(after.json()["config"]["spacing"]["line_spacing"], 1.6)

        tab_update = self.client.post(
            "/api/config/update",
            json={"path": "spacing.tab_width", "value": 8},
        )
        self.assertEqual(tab_update.status_code, 200)
        after_tab = self.client.get("/api/config")
        self.assertEqual(after_tab.status_code, 200)
        self.assertEqual(after_tab.json()["config"]["spacing"]["tab_width"], 8)

    def test_apply_theme_updates_config(self) -> None:
        applied = self.client.post(
            "/api/themes/apply",
            json={"theme_name": "corporate"},
        )
        self.assertEqual(applied.status_code, 200)
        payload = applied.json()
        self.assertTrue(payload.get("success"))
        cfg = payload.get("config", {})
        self.assertEqual(cfg.get("app", {}).get("theme"), "corporate")
        self.assertEqual(cfg.get("fonts", {}).get("family"), "Arial")
        self.assertTrue(cfg.get("colors", {}).get("h1"))

    @patch("app.exporter._convert_docx_to_pdf", return_value=(False, "converter unavailable"))
    @patch("app.exporter._convert_html_to_pdf_weasyprint", return_value=(False, "weasy unavailable"))
    @patch("app.exporter._convert_nodes_to_pdf_reportlab", return_value=(False, "reportlab unavailable"))
    def test_generate_pdf_fallback_contract(self, _mock_reportlab, _mock_weasy, _mock_docx) -> None:
        response = self.client.post(
            "/api/generate",
            json={
                "content": SAMPLE_EXAMPLE,
                "theme": PROFESSIONAL_THEME.model_dump(),
                "format": "pdf",
                "filename": "notesforge_output",
                "security": {"disableEditingDocx": False, "removeMetadata": True},
                "templateId": "assignment",
            },
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload.get("requestedFormat"), "pdf")
        self.assertEqual(payload.get("actualFormat"), "pdf")
        self.assertIn(payload.get("conversionEngine"), {"weasyprint", "reportlab", "emergency_pdf", "ilovepdf_api", "ilovepdf_automation"})
        joined_warnings = " | ".join(payload.get("warnings") or [])
        self.assertIn("fallback renderer", joined_warnings.lower())
        self.assertTrue(payload.get("downloadUrl", "").startswith("/api/download/"))
        dl = self.client.get(payload["downloadUrl"])
        self.assertEqual(dl.status_code, 200)
        self.assertTrue(dl.headers["content-type"].startswith("application/pdf"))

    def test_generate_async_job_lifecycle(self) -> None:
        response = self.client.post(
            "/api/generate/async",
            json={
                "content": SAMPLE_EXAMPLE,
                "theme": PROFESSIONAL_THEME.model_dump(),
                "format": "docx",
                "filename": "async_output",
                "security": {"disableEditingDocx": False, "removeMetadata": False},
            },
        )
        self.assertEqual(response.status_code, 202)
        body = response.json()
        self.assertTrue(body.get("success"))
        job_id = body.get("jobId")
        self.assertTrue(isinstance(job_id, str) and len(job_id) >= 20)

        final = None
        for _ in range(80):
            poll = self.client.get(f"/api/generate/jobs/{job_id}")
            self.assertEqual(poll.status_code, 200)
            payload = poll.json()
            if payload.get("status") in {"completed", "failed"}:
                final = payload
                break
            time.sleep(0.05)

        self.assertIsNotNone(final)
        self.assertEqual(final.get("status"), "completed")
        self.assertEqual(final.get("requestedFormat"), "docx")
        self.assertEqual(final.get("actualFormat"), "docx")
        self.assertTrue(final.get("downloadUrl", "").startswith("/api/download/"))

        direct_download = self.client.get(f"/api/generate/jobs/{job_id}/download")
        self.assertEqual(direct_download.status_code, 200)
        self.assertTrue(
            direct_download.headers["content-type"].startswith(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        )

    def test_file_processing_context_contract(self) -> None:
        response = self.client.get("/api/file-processing/context")
        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertIn("runtimeTarget", body)
        self.assertIn("platform", body)
        self.assertIn("directories", body)
        self.assertIn("markerTemplateExample", body)
        self.assertIn("outputDirectory", body)
        self.assertTrue(isinstance(body["directories"], list))

    @patch("app.pdf_conversion._convert_pdf_to_docx_local")
    def test_file_processing_convert_pdf_to_docx(self, mock_convert) -> None:
        def fake_convert(source_path, output_path):
            document = Document()
            document.add_paragraph("Converted from pdf")
            document.save(str(output_path))
            return True, "pdf2docx"

        mock_convert.side_effect = fake_convert

        response = self.client.post(
            "/api/file-processing/convert",
            data={
                "target_format": "docx",
                "provider_preference": "local",
                "preserve_layout": "true",
            },
            files={"file": ("sample.pdf", self._sample_pdf_bytes(), "application/pdf")},
        )
        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body.get("success"))
        self.assertEqual(body.get("sourceFormat"), "pdf")
        self.assertEqual(body.get("targetFormat"), "docx")
        self.assertEqual(body.get("actualFormat"), "docx")
        self.assertEqual(body.get("providerUsed"), "local")
        self.assertEqual(body.get("conversionEngine"), "pdf2docx")
        download = self.client.get(body["downloadUrl"])
        self.assertEqual(download.status_code, 200)
        self.assertTrue(
            download.headers["content-type"].startswith(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        )

    @patch("app.pdf_conversion._convert_docx_to_pdf_local")
    def test_file_processing_convert_docx_to_pdf(self, mock_convert) -> None:
        def fake_convert(source_path, output_path):
            writer = PdfWriter()
            writer.add_blank_page(width=300, height=300)
            with open(output_path, "wb") as stream:
                writer.write(stream)
            return True, "docx2pdf"

        mock_convert.side_effect = fake_convert

        response = self.client.post(
            "/api/file-processing/convert",
            data={
                "target_format": "pdf",
                "provider_preference": "local",
                "preserve_layout": "true",
            },
            files={
                "file": (
                    "sample.docx",
                    self._sample_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body.get("success"))
        self.assertEqual(body.get("sourceFormat"), "docx")
        self.assertEqual(body.get("targetFormat"), "pdf")
        self.assertEqual(body.get("actualFormat"), "pdf")
        self.assertEqual(body.get("providerUsed"), "local")
        self.assertEqual(body.get("conversionEngine"), "docx2pdf")
        download = self.client.get(body["downloadUrl"])
        self.assertEqual(download.status_code, 200)
        self.assertTrue(download.headers["content-type"].startswith("application/pdf"))

    @patch("app.pdf_conversion._convert_docx_to_pdf_local")
    def test_file_processing_batch_convert_returns_zip(self, mock_convert) -> None:
        def fake_convert(source_path, output_path):
            writer = PdfWriter()
            writer.add_blank_page(width=300, height=300)
            with open(output_path, "wb") as stream:
                writer.write(stream)
            return True, "docx2pdf"

        mock_convert.side_effect = fake_convert

        response = self.client.post(
            "/api/file-processing/convert/batch",
            data={
                "target_format": "pdf",
                "provider_preference": "local",
                "preserve_layout": "true",
            },
            files=[
                (
                    "files",
                    (
                        "batch_one.docx",
                        self._sample_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "files",
                    (
                        "batch_two.docx",
                        self._sample_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        self.assertEqual(response.status_code, 202)
        job_id = response.json()["jobId"]

        final = None
        for _ in range(80):
            poll = self.client.get(f"/api/file-processing/convert/batch/jobs/{job_id}")
            self.assertEqual(poll.status_code, 200)
            payload = poll.json()
            if payload.get("status") in {"completed", "failed"}:
                final = payload
                break
            time.sleep(0.05)

        self.assertIsNotNone(final)
        self.assertEqual(final.get("status"), "completed")
        self.assertEqual(final.get("actualFormat"), "zip")
        self.assertTrue(final.get("downloadUrl", "").startswith("/api/download/"))

        download = self.client.get(final["downloadUrl"])
        self.assertEqual(download.status_code, 200)
        self.assertTrue(download.headers["content-type"].startswith("application/zip"))
        with zipfile.ZipFile(io.BytesIO(download.content)) as archive:
            names = set(archive.namelist())
        self.assertIn("batch_one_converted.pdf", names)
        self.assertIn("batch_two_converted.pdf", names)
        self.assertIn("batch_report.json", names)

    def test_download_rejects_non_token_path(self) -> None:
        response = self.client.get("/api/download/../../etc/passwd")
        self.assertIn(response.status_code, (404, 422))


if __name__ == "__main__":
    unittest.main()
