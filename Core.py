"""
Core.py - Core Logic for NotesForge
Contains all text analysis and document building logic (PROVEN v3 code!)
Supports both DOCX and PDF output
"""

import re
import tempfile
from typing import List, Tuple, Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class TextAnalyzer:
    """
    Analyzes and classifies text (PROVEN v3 logic!)
    Detects headings, lists, code, tables, formulas, etc.
    """
    
    def __init__(self, commands_list):
        self.commands = commands_list
    
    def clean_line(self, line: str) -> str:
        """Clean markdown and formatting"""
        line = line.strip()
        
        # Remove markdown
        line = line.replace("**", "")
        line = line.replace("```bash", "")
        line = line.replace("```python", "")
        line = line.replace("```javascript", "")
        line = line.replace("```", "")
        
        # Remove separators
        if line in ["•", "--", "---", "___", "***", "─", "=="]:
            return ""
        
        # Remove markdown headings
        line = re.sub(r"^#+\s*", "", line)
        
        return line.strip()
    
    def classify_line(self, line: str, next_line: str = "", prev_line: str = "") -> Dict:
        """
        CORE CLASSIFICATION FUNCTION (PROVEN FROM v3!)
        Returns: {'type': str, 'content': str, 'metadata': dict}
        """
        
        line_clean = self.clean_line(line)
        
        # Empty line
        if not line_clean:
            return {'type': 'empty', 'content': '', 'metadata': {}}
        
        # === HEADINGS ===
        
        # Day heading: "Day 23 -" or "Day 23 - Title"
        day_match = re.match(r'^Day\s+(\d+)\s*[-–—:]?\s*(.*)$', line_clean, re.IGNORECASE)
        if day_match:
            day_num = day_match.group(1)
            title = day_match.group(2).strip()
            
            if title:
                return {
                    'type': 'heading_day_with_title',
                    'content': f"Day {day_num} - {title}",
                    'metadata': {'day': day_num, 'title': title}
                }
            else:
                return {
                    'type': 'heading_day',
                    'content': f"Day {day_num} -",
                    'metadata': {'day': day_num}
                }
        
        # Topic: "1. Topic: Name"
        if re.match(r'^\d+\.\s+Topic:', line_clean, re.IGNORECASE):
            return {'type': 'topic_heading', 'content': line_clean, 'metadata': {}}
        
        # Section: "1. Name" (no Topic:)
        if re.match(r'^\d+\.\s+\w', line_clean) and 'topic:' not in line_clean.lower():
            return {'type': 'section_heading', 'content': line_clean, 'metadata': {}}
        
        # === LABELS ===
        
        if re.match(r'^(Example|Examples|Step|Steps|Note|Notes|Solution|Answer|Output|Input|Result|Commands?)s?:', 
                   line_clean, re.IGNORECASE):
            return {'type': 'label', 'content': line_clean, 'metadata': {}}
        
        # === LISTS ===
        
        # Bullet points
        if re.match(r'^[•\-*]\s+', line_clean):
            content = re.sub(r'^[•\-*]\s+', '', line_clean)
            return {'type': 'bullet', 'content': content, 'metadata': {}}
        
        # Numbered lists
        if re.match(r'^\d+[)\-]\s+\w', line_clean):
            content = re.sub(r'^\d+[)\-]\s+', '', line_clean)
            return {'type': 'numbered_list', 'content': content, 'metadata': {}}
        
        # === CODE & COMMANDS ===
        
        # Command line prefix
        if re.match(r'^[$#]\s+', line_clean):
            return {'type': 'command', 'content': line_clean, 'metadata': {}}
        
        # Known commands
        for cmd in self.commands:
            if re.match(rf'^{cmd}\b', line_clean):
                return {'type': 'command', 'content': line_clean, 'metadata': {'command': cmd}}
        
        # File paths
        if re.search(r'[/\\][\w\-/.]+', line_clean) and len(line_clean) < 100:
            return {'type': 'filepath', 'content': line_clean, 'metadata': {}}
        
        # Code patterns (more strict to avoid false positives)
        # Only detect as code if it has MULTIPLE code indicators
        code_indicators = 0
        if re.search(r'[\{\}]', line_clean):  # Curly braces
            code_indicators += 1
        if re.search(r'[\[\]]', line_clean) and not re.search(r'\w+\s+\[', line_clean):  # Brackets (but not "word [")
            code_indicators += 1
        if re.search(r'\w+\s*\(.*\)\s*[\{;]', line_clean):  # function() { or function();
            code_indicators += 1
        if re.search(r'^(def|class|function|var|let|const|public|private)\s+', line_clean):  # Code keywords
            code_indicators += 1
        if re.search(r'[;{}]\s*$', line_clean):  # Ends with ; or } or {
            code_indicators += 1
        
        # Only classify as code if 2+ indicators (more strict)
        if code_indicators >= 2 and not re.search(r'\d+\s*[×÷=]\s*\d+', line_clean):
            return {'type': 'code', 'content': line_clean, 'metadata': {}}
        
        # === SPECIAL CONTENT ===
        
        # Quotes
        if (line_clean.startswith('"') and line_clean.endswith('"')) or \
           (line_clean.startswith("'") and line_clean.endswith("'")):
            return {'type': 'quote', 'content': line_clean, 'metadata': {}}
        
        # Formulas
        if re.search(r'[×÷=\^²³¹⁰₀₁₂₃]|x\s*10|10\^|\d+\s*[+\-*/]\s*\d+', line_clean):
            return {'type': 'formula', 'content': line_clean, 'metadata': {}}
        
        # Table row
        if self.is_table_row(line_clean, line):
            return {'type': 'table_row', 'content': line_clean, 'metadata': {}}
        
        # === DEFAULT: PARAGRAPH ===
        
        return {'type': 'paragraph', 'content': line_clean, 'metadata': {}}
    
    def is_table_row(self, line_clean: str, line_raw: str) -> bool:
        """Detect table rows"""
        
        # Pipe separators
        if line_clean.count('|') >= 2:
            return True
        
        # Space separators (3+)
        if re.search(r'\S\s{3,}\S', line_raw):
            parts = re.split(r'\s{3,}', line_clean.strip())
            if len(parts) >= 2:
                return True
        
        # Tab separators
        if '\t' in line_raw:
            parts = line_raw.split('\t')
            if len([p for p in parts if p.strip()]) >= 2:
                return True
        
        return False
    
    def parse_table(self, lines: List[str], start_idx: int) -> Tuple[int, List[List[str]]]:
        """Parse table rows"""
        
        table_rows = []
        idx = start_idx
        
        while idx < len(lines):
            line = lines[idx].strip()
            
            if not line or not self.is_table_row(line, lines[idx]):
                break
            
            # Skip separator rows
            if re.match(r'^[-–—=|:\s]+$', line):
                idx += 1
                continue
            
            # Parse columns
            if '|' in line:
                cols = [c.strip() for c in line.split('|')]
                cols = [c for c in cols if c]
            elif '\t' in lines[idx]:
                cols = [c.strip() for c in lines[idx].split('\t')]
                cols = [c for c in cols if c]
            else:
                cols = [c.strip() for c in re.split(r'\s{3,}', line)]
            
            if cols:
                table_rows.append(cols)
            
            idx += 1
        
        return idx - 1, table_rows
    
    def analyze_document(self, text: str) -> Dict[str, int]:
        """Analyze document and return statistics"""
        
        lines = text.split('\n')
        stats = {
            'total_lines': len([l for l in lines if l.strip()]),
            'day_headings': 0,
            'topics': 0,
            'sections': 0,
            'bullets': 0,
            'commands': 0,
            'code_blocks': 0,
            'tables': 0,
            'labels': 0,
            'formulas': 0
        }
        
        in_table = False
        
        for line in lines:
            c = self.classify_line(line)
            ctype = c['type']
            
            if 'heading_day' in ctype:
                stats['day_headings'] += 1
            elif ctype == 'topic_heading':
                stats['topics'] += 1
            elif ctype == 'section_heading':
                stats['sections'] += 1
            elif ctype == 'bullet':
                stats['bullets'] += 1
            elif ctype == 'command':
                stats['commands'] += 1
            elif ctype in ['code', 'filepath']:
                stats['code_blocks'] += 1
            elif ctype == 'table_row':
                if not in_table:
                    stats['tables'] += 1
                    in_table = True
            elif ctype == 'label':
                stats['labels'] += 1
            elif ctype == 'formula':
                stats['formulas'] += 1
            else:
                in_table = False
        
        return stats


class DocumentBuilder:
    """
    Builds Word documents (PROVEN v3 logic!)
    Creates professional DOCX files with all styling
    """
    
    def __init__(self, config, analyzer):
        self.config = config
        self.analyzer = analyzer
        self.doc = Document()
    
    def add_page_border(self):
        """Add 4-side page border"""
        if not self.config.get('page.border.enabled', True):
            return
        
        section = self.doc.sections[0]
        sectPr = section._sectPr
        
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(self.config.get('page.border.width', 12)))
            border.set(qn('w:space'), str(self.config.get('page.border.space', 24)))
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
        
        sectPr.append(pgBorders)
    
    def add_header_footer(self):
        """Add header and footer"""
        section = self.doc.sections[0]
        
        # Header
        if self.config.get('header.enabled', True):
            header = section.header.paragraphs[0]
            header.text = self.config.get('header.text', 'NotesForge')
            header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            run = header.runs[0]
            run.font.size = Pt(self.config.get('header.size', 11))
            run.font.bold = self.config.get('header.bold', True)
            
            color_hex = self.config.get('header.color', '#FF8C00')
            rgb = self.config.hex_to_rgb(color_hex)
            run.font.color.rgb = RGBColor(*rgb)
            run.font.name = self.config.get('fonts.family', 'Times New Roman')
        
        # Footer
        if self.config.get('footer.enabled', True):
            footer = section.footer.paragraphs[0]
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if self.config.get('footer.show_page_numbers', True):
                run = footer.add_run("Page ")
                run.font.size = Pt(self.config.get('footer.size', 10))
                run.font.name = self.config.get('fonts.family', 'Times New Roman')
                
                # Page number field
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = "PAGE"
                
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
    
    def add_paragraph(self, text: str, style_type: str = 'paragraph'):
        """Add styled paragraph"""
        
        # Create paragraph
        if style_type == 'bullet':
            p = self.doc.add_paragraph(style='List Bullet')
        elif style_type == 'numbered_list':
            p = self.doc.add_paragraph(style='List Number')
        else:
            p = self.doc.add_paragraph()
        
        # Base spacing
        p.paragraph_format.line_spacing = self.config.get('spacing.line_spacing', 1.5)
        p.paragraph_format.space_after = Pt(self.config.get('spacing.paragraph_spacing_after', 6))
        
        # Font family
        font_family = self.config.get('fonts.family', 'Times New Roman')
        
        # Style-specific formatting
        if style_type == 'heading_day':
            p.paragraph_format.space_before = Pt(self.config.get('spacing.heading_spacing_before', 12))
            p.paragraph_format.space_after = Pt(self.config.get('spacing.heading_spacing_after', 6))
            run = p.add_run(text)
            run.font.name = font_family
            run.font.size = Pt(self.config.get('fonts.sizes.day_heading', 18))
            run.bold = True
            rgb = self.config.hex_to_rgb(self.config.get('colors.day_heading', '#FF8C00'))
            run.font.color.rgb = RGBColor(*rgb)
            
        elif style_type == 'heading_day_with_title':
            p.paragraph_format.space_before = Pt(self.config.get('spacing.heading_spacing_before', 12))
            p.paragraph_format.space_after = Pt(self.config.get('spacing.heading_spacing_after', 6))
            run = p.add_run(text)
            run.font.name = font_family
            run.font.size = Pt(self.config.get('fonts.sizes.day_heading_with_title', 16))
            run.bold = True
            rgb = self.config.hex_to_rgb(self.config.get('colors.day_heading', '#FF8C00'))
            run.font.color.rgb = RGBColor(*rgb)
            
        elif style_type == 'topic_heading':
            p.paragraph_format.space_before = Pt(self.config.get('spacing.heading_spacing_before', 12))
            p.paragraph_format.space_after = Pt(self.config.get('spacing.heading_spacing_after', 6))
            run = p.add_run(text)
            run.font.name = font_family
            run.font.size = Pt(self.config.get('fonts.sizes.topic_heading', 14))
            run.bold = True
            rgb = self.config.hex_to_rgb(self.config.get('colors.topic_heading', '#1F4788'))
            run.font.color.rgb = RGBColor(*rgb)
            
        elif style_type == 'section_heading':
            p.paragraph_format.space_before = Pt(self.config.get('spacing.heading_spacing_before', 12))
            p.paragraph_format.space_after = Pt(self.config.get('spacing.heading_spacing_after', 6))
            run = p.add_run(text)
            run.font.name = font_family
            run.font.size = Pt(self.config.get('fonts.sizes.section_heading', 13))
            run.bold = True
            rgb = self.config.hex_to_rgb(self.config.get('colors.section_heading', '#2E5C8A'))
            run.font.color.rgb = RGBColor(*rgb)
            
        elif style_type == 'label':
            run = p.add_run(text)
            run.font.name = font_family
            run.font.size = Pt(self.config.get('fonts.sizes.label', 12))
            run.bold = True
            
        elif style_type in ['command', 'code', 'filepath']:
            run = p.add_run(text)
            run.font.name = self.config.get('fonts.family_code', 'Courier New')
            run.font.size = Pt(self.config.get('fonts.sizes.code', 10))
            
            # Gray background
            bg_color = self.config.get('colors.code_background', '#F5F5F5').lstrip('#')
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg_color)
            p._p.get_or_add_pPr().append(shading)
            
            p.paragraph_format.left_indent = Inches(self.config.get('spacing.code_indent', 0.25))
            
        elif style_type == 'quote':
            run = p.add_run(text)
            run.font.name = font_family
            run.font.size = Pt(self.config.get('fonts.sizes.quote', 12))
            run.italic = True
            p.paragraph_format.left_indent = Inches(self.config.get('spacing.quote_indent', 0.5))
            
        elif style_type == 'formula':
            run = p.add_run(text)
            run.font.name = font_family
            run.font.size = Pt(self.config.get('fonts.sizes.formula', 12))
            
        else:
            # Regular paragraph
            run = p.add_run(text)
            run.font.name = font_family
            run.font.size = Pt(self.config.get('fonts.sizes.body', 12))
    
    def add_table(self, table_data: List[List[str]]):
        """Create professional table"""
        
        if not table_data or len(table_data) < 1:
            return
        
        # Determine columns
        num_cols = max(len(row) for row in table_data)
        
        # Create table
        table = self.doc.add_table(rows=len(table_data), cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Fill table
        for row_idx, row_data in enumerate(table_data):
            row_cells = table.rows[row_idx].cells
            
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row_cells):
                    cell = row_cells[col_idx]
                    cell.text = cell_text
                    
                    # Style text
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = self.config.get('fonts.family', 'Times New Roman')
                            run.font.size = Pt(11)
                            
                            # Bold header
                            if row_idx == 0:
                                run.bold = True
                                header_color = self.config.get('colors.table_header_text', '#FFFFFF')
                                rgb = self.config.hex_to_rgb(header_color)
                                run.font.color.rgb = RGBColor(*rgb)
                    
                    # Cell shading
                    shading = OxmlElement('w:shd')
                    if row_idx == 0:
                        color = self.config.get('colors.table_header_bg', '#4A7BA7').lstrip('#')
                        shading.set(qn('w:fill'), color)
                    elif row_idx % 2 == 1:
                        color = self.config.get('colors.table_odd_row', '#F5F5F5').lstrip('#')
                        shading.set(qn('w:fill'), color)
                    else:
                        color = self.config.get('colors.table_even_row', '#FFFFFF').lstrip('#')
                        shading.set(qn('w:fill'), color)
                    
                    cell._element.get_or_add_tcPr().append(shading)
    
    def build_from_text(self, text: str) -> str:
        """Main function: Build document from text"""
        
        # Setup document
        self.add_header_footer()
        self.add_page_border()
        
        # Set margins
        section = self.doc.sections[0]
        section.top_margin = Inches(self.config.get('page.margins.top', 1.0))
        section.bottom_margin = Inches(self.config.get('page.margins.bottom', 1.0))
        section.left_margin = Inches(self.config.get('page.margins.left', 1.0))
        section.right_margin = Inches(self.config.get('page.margins.right', 1.0))
        
        # Parse text
        lines = text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            next_line = lines[i + 1] if i + 1 < len(lines) else ""
            prev_line = lines[i - 1] if i > 0 else ""
            
            # Classify line
            classification = self.analyzer.classify_line(line, next_line, prev_line)
            line_type = classification['type']
            content = classification['content']
            
            # Skip empty
            if line_type == 'empty':
                i += 1
                continue
            
            # Handle tables
            if line_type == 'table_row':
                end_idx, table_data = self.analyzer.parse_table(lines, i)
                self.add_table(table_data)
                i = end_idx + 1
                continue
            
            # Add content
            if line_type in ['heading_day', 'heading_day_with_title', 'topic_heading', 
                           'section_heading', 'label', 'command', 'code', 'filepath',
                           'quote', 'formula']:
                self.add_paragraph(content, style_type=line_type)
            elif line_type == 'bullet':
                self.add_paragraph(content, style_type='bullet')
            elif line_type == 'numbered_list':
                self.add_paragraph(content, style_type='numbered_list')
            else:
                self.add_paragraph(content, style_type='paragraph')
            
            i += 1
        
        # Save to temp file
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        self.doc.save(tmp.name)
        return tmp.name


class PDFBuilder:
    """
    Builds PDF documents
    Uses docx2pdf with proper COM initialization on Windows
    Falls back to DOCX if PDF conversion fails
    """
    
    def __init__(self, config, analyzer):
        self.config = config
        self.analyzer = analyzer
    
    def build_from_text(self, text: str) -> str:
        """Build PDF from text"""
        # First create DOCX
        docx_builder = DocumentBuilder(self.config, self.analyzer)
        docx_path = docx_builder.build_from_text(text)
        
        # Try to convert to PDF
        try:
            import sys
            import os
            
            # Check if running on Windows
            if sys.platform == 'win32':
                try:
                    # Try docx2pdf with COM initialization fix
                    import pythoncom
                    pythoncom.CoInitialize()
                    
                    from docx2pdf import convert
                    pdf_path = docx_path.replace('.docx', '.pdf')
                    convert(docx_path, pdf_path)
                    
                    pythoncom.CoUninitialize()
                    
                    if os.path.exists(pdf_path):
                        return pdf_path
                    else:
                        raise Exception("PDF file was not created")
                        
                except Exception as e:
                    # COM initialization failed or Word not installed
                    # Return DOCX instead
                    import warnings
                    warnings.warn(f"PDF conversion failed: {str(e)}. Returning DOCX instead.")
                    return docx_path
            else:
                # On Linux/Mac, try different approach
                try:
                    from docx2pdf import convert
                    pdf_path = docx_path.replace('.docx', '.pdf')
                    convert(docx_path, pdf_path)
                    return pdf_path
                except:
                    return docx_path
                    
        except ImportError:
            # docx2pdf not installed, return DOCX
            return docx_path
